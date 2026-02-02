#!/usr/bin/env python3
"""
TechWriterReview - Comment Inserter Module
==========================================
Version: reads from version.json (module v2.1)

Inserts DOCX comments at broken hyperlink locations to notify document authors
of link issues that need attention.

Features:
- Insert comments at broken/problematic hyperlinks in DOCX files
- Comprehensive hyperlink discovery: paragraphs, tables, headers/footers
- Generate comment pack with location hints for manual application
- Configurable comment scope (broken_only vs all_non_working)
- Preserve original document formatting
- v3.0.109: Improved text matching with smart quote normalization and multi-strategy fallback

v3.0.38 Batch G: Initial implementation for hyperlink comments.
v3.0.39 Batch G-fix: Comprehensive hyperlink discovery, location hints,
                     comment scope filtering, URL redaction in logs.
v3.0.109: Issue #11 - Comment placement improved text matching:
          - Smart quotes normalization (", ", ', ' → ", ')
          - Whitespace normalization (collapse multiple spaces/tabs/newlines)
          - Multi-strategy matching: exact → normalized → fuzzy
          - Enhanced logging for debugging which strategy succeeded
"""

import os
import shutil
import tempfile
import logging
from pathlib import Path
from datetime import datetime
from typing import List, Dict, Optional, Tuple, Set
from dataclasses import dataclass, field
from enum import Enum
import re  # v3.0.109: Added for text normalization regex operations

__version__ = "2.1.0"  # v3.0.109: Added improved text matching

# Set up module logger
_logger = logging.getLogger('comment_inserter')


# =============================================================================
# v3.0.109: TEXT NORMALIZATION AND MATCHING UTILITIES
# =============================================================================
# These functions address Issue #11: Comment placement improved text matching
# Problem: Comments sometimes attach to wrong text due to:
#   1. Smart quotes vs straight quotes mismatch
#   2. Whitespace normalization differences
#   3. Text spanning multiple XML runs
# =============================================================================

def normalize_quotes(text: str) -> str:
    """
    v3.0.109: Normalize smart/curly quotes to straight quotes for matching.
    
    Converts:
    - Left/right double quotes (", ") → straight double quote (")
    - Left/right single quotes (', ') → straight single quote (')
    - Various other quote characters used in Word documents
    
    Args:
        text: Input text that may contain smart quotes
        
    Returns:
        Text with all quotes normalized to straight ASCII quotes
    """
    if not text:
        return ""
    
    # Smart/curly double quotes → straight double quote
    text = text.replace('\u201C', '"')   # Left double quote U+201C "
    text = text.replace('\u201D', '"')   # Right double quote U+201D "
    text = text.replace('\u201E', '"')   # Double low-9 quote U+201E „
    text = text.replace('\u201F', '"')   # Double high-reversed-9 quote U+201F ‟
    text = text.replace('\u00AB', '"')   # Left-pointing double angle quote U+00AB «
    text = text.replace('\u00BB', '"')   # Right-pointing double angle quote U+00BB »
    
    # Smart/curly single quotes → straight single quote
    text = text.replace('\u2018', "'")   # Left single quote U+2018 '
    text = text.replace('\u2019', "'")   # Right single quote U+2019 '
    text = text.replace('\u201A', "'")   # Single low-9 quote U+201A ‚
    text = text.replace('\u201B', "'")   # Single high-reversed-9 quote U+201B ‛
    text = text.replace('\u2039', "'")   # Single left-pointing angle quote U+2039 ‹
    text = text.replace('\u203A', "'")   # Single right-pointing angle quote U+203A ›
    text = text.replace('\u0060', "'")   # Grave accent U+0060 `
    text = text.replace('\u00B4', "'")   # Acute accent U+00B4 ´
    
    return text


def normalize_whitespace(text: str) -> str:
    """
    v3.0.109: Normalize whitespace for consistent text matching.
    
    Converts:
    - Multiple spaces → single space
    - Tabs → single space
    - Newlines (\\n, \\r) → single space
    - Non-breaking spaces → regular space
    - Other Unicode whitespace → regular space
    
    Args:
        text: Input text with potentially irregular whitespace
        
    Returns:
        Text with normalized whitespace (single spaces, trimmed)
    """
    if not text:
        return ""
    
    # Replace non-breaking spaces and other Unicode spaces with regular space
    text = text.replace('\u00A0', ' ')   # Non-breaking space
    text = text.replace('\u2007', ' ')   # Figure space
    text = text.replace('\u2008', ' ')   # Punctuation space
    text = text.replace('\u2009', ' ')   # Thin space
    text = text.replace('\u200A', ' ')   # Hair space
    text = text.replace('\u200B', '')    # Zero-width space (remove)
    text = text.replace('\u202F', ' ')   # Narrow no-break space
    text = text.replace('\u205F', ' ')   # Medium mathematical space
    text = text.replace('\u3000', ' ')   # Ideographic space
    
    # Replace tabs, newlines with space
    text = text.replace('\t', ' ')
    text = text.replace('\r\n', ' ')
    text = text.replace('\r', ' ')
    text = text.replace('\n', ' ')
    
    # Collapse multiple spaces to single space using regex
    text = re.sub(r'\s+', ' ', text)
    
    # Strip leading/trailing whitespace
    return text.strip()


def normalize_text_for_matching(text: str) -> str:
    """
    v3.0.109: Full text normalization combining quotes and whitespace.
    
    Use this function when preparing text for comparison/matching.
    
    Args:
        text: Input text to normalize
        
    Returns:
        Fully normalized text ready for matching
    """
    if not text:
        return ""
    
    text = normalize_quotes(text)
    text = normalize_whitespace(text)
    return text


class TextMatchResult:
    """
    v3.0.109: Result of a text matching operation.
    
    Tracks which matching strategy succeeded and provides debug info.
    """
    def __init__(self, found: bool, strategy: str, position: int = -1, 
                 matched_text: str = "", search_text: str = ""):
        self.found = found
        self.strategy = strategy      # 'exact', 'normalized', 'fuzzy', 'none'
        self.position = position      # Character position in document (-1 if not found)
        self.matched_text = matched_text  # Actual text that matched
        self.search_text = search_text    # Original search text
    
    def __bool__(self):
        return self.found
    
    def __repr__(self):
        if self.found:
            return f"TextMatchResult(found=True, strategy='{self.strategy}', pos={self.position})"
        return f"TextMatchResult(found=False, strategy='none')"


def find_text_in_document(
    search_text: str,
    document_text: str,
    fuzzy_min_length: int = 20,
    fuzzy_char_count: int = 30,
    enable_logging: bool = True
) -> TextMatchResult:
    """
    v3.0.109: Find text in document using multiple matching strategies.
    
    Tries matching strategies in order of strictness:
    1. EXACT: Direct substring match
    2. NORMALIZED: Both texts normalized (quotes + whitespace)
    3. FUZZY: Match first N characters (for long text that may be truncated)
    
    Args:
        search_text: The text to find
        document_text: The document content to search in
        fuzzy_min_length: Minimum search_text length to try fuzzy match (default 20)
        fuzzy_char_count: Number of characters to match in fuzzy mode (default 30)
        enable_logging: Whether to log the matching results
        
    Returns:
        TextMatchResult with match details and strategy used
    """
    if not search_text or not document_text:
        if enable_logging:
            _logger.debug("[TextMatch] Empty search_text or document_text")
        return TextMatchResult(found=False, strategy='none', search_text=search_text)
    
    # Strategy 1: EXACT match
    if search_text in document_text:
        pos = document_text.find(search_text)
        if enable_logging:
            _logger.debug(f"[TextMatch] EXACT match found at position {pos}: '{search_text[:40]}...'")
        return TextMatchResult(
            found=True, 
            strategy='exact', 
            position=pos,
            matched_text=search_text,
            search_text=search_text
        )
    
    # Strategy 2: NORMALIZED match (quotes + whitespace)
    norm_search = normalize_text_for_matching(search_text)
    norm_document = normalize_text_for_matching(document_text)
    
    if norm_search and norm_search in norm_document:
        pos = norm_document.find(norm_search)
        if enable_logging:
            _logger.debug(f"[TextMatch] NORMALIZED match found at position {pos}: '{search_text[:40]}...'")
        return TextMatchResult(
            found=True, 
            strategy='normalized', 
            position=pos,
            matched_text=norm_search,
            search_text=search_text
        )
    
    # Strategy 3: FUZZY match (first N characters)
    # Only try this for longer search texts to avoid false positives
    if len(norm_search) >= fuzzy_min_length:
        # Take first N characters (adjusting if search is shorter)
        chars_to_match = min(fuzzy_char_count, len(norm_search))
        fuzzy_search = norm_search[:chars_to_match]
        
        if fuzzy_search in norm_document:
            pos = norm_document.find(fuzzy_search)
            if enable_logging:
                _logger.debug(f"[TextMatch] FUZZY match (first {chars_to_match} chars) at position {pos}: '{search_text[:40]}...'")
            return TextMatchResult(
                found=True, 
                strategy='fuzzy', 
                position=pos,
                matched_text=fuzzy_search,
                search_text=search_text
            )
    
    # No match found
    if enable_logging:
        _logger.debug(f"[TextMatch] No match found for: '{search_text[:60]}...'")
        # Log what we tried for debugging
        _logger.debug(f"[TextMatch] Original search ({len(search_text)} chars): '{search_text[:80]}'")
        _logger.debug(f"[TextMatch] Normalized search ({len(norm_search)} chars): '{norm_search[:80]}'")
    
    return TextMatchResult(found=False, strategy='none', search_text=search_text)


def find_text_in_xml_runs(
    search_text: str,
    paragraph_element,
    fuzzy_min_length: int = 20,
    fuzzy_char_count: int = 30,
    enable_logging: bool = True
) -> TextMatchResult:
    """
    v3.0.109: Find text across multiple XML runs in a Word paragraph element.
    
    Word documents split text across multiple <w:r> runs due to formatting changes,
    spell-check, etc. This function reconstructs the full paragraph text and finds
    the search text across run boundaries.
    
    Args:
        search_text: The text to find
        paragraph_element: lxml element representing a Word paragraph (<w:p>)
        fuzzy_min_length: Minimum search_text length to try fuzzy match
        fuzzy_char_count: Number of characters to match in fuzzy mode
        enable_logging: Whether to log the matching results
        
    Returns:
        TextMatchResult with match details
    """
    # Define Word namespace
    W_NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
    
    try:
        # Collect all text from <w:t> elements within the paragraph
        text_elements = paragraph_element.findall(f'.//{W_NS}t')
        
        if not text_elements:
            # Try without namespace prefix (for compatibility)
            text_elements = paragraph_element.findall('.//t')
        
        # Reconstruct full paragraph text
        para_text_parts = []
        for t_elem in text_elements:
            if t_elem.text:
                para_text_parts.append(t_elem.text)
        
        para_text = ''.join(para_text_parts)
        
        # Use the main find function
        return find_text_in_document(
            search_text, 
            para_text,
            fuzzy_min_length=fuzzy_min_length,
            fuzzy_char_count=fuzzy_char_count,
            enable_logging=enable_logging
        )
        
    except Exception as e:
        if enable_logging:
            _logger.warning(f"[TextMatch] Error searching XML runs: {e}")
        return TextMatchResult(found=False, strategy='none', search_text=search_text)


def find_text_in_docx_paragraphs(
    search_text: str,
    paragraphs: list,
    fuzzy_min_length: int = 20,
    fuzzy_char_count: int = 30,
    enable_logging: bool = True
) -> tuple:
    """
    v3.0.109: Find text across multiple paragraphs in a DOCX document.
    
    Searches through a list of paragraph elements and returns the first
    matching paragraph along with match details.
    
    Args:
        search_text: The text to find
        paragraphs: List of lxml paragraph elements
        fuzzy_min_length: Minimum search_text length to try fuzzy match
        fuzzy_char_count: Number of characters to match in fuzzy mode
        enable_logging: Whether to log the matching results
        
    Returns:
        Tuple of (paragraph_element, TextMatchResult) or (None, TextMatchResult) if not found
    """
    W_NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
    
    for para_idx, para in enumerate(paragraphs):
        try:
            # Get paragraph text - try with namespace first
            text_elements = para.findall(f'.//{W_NS}t')
            if not text_elements:
                text_elements = para.findall('.//t')
            
            para_text = ''.join(t.text or '' for t in text_elements)
            
            # Try to match
            result = find_text_in_document(
                search_text,
                para_text,
                fuzzy_min_length=fuzzy_min_length,
                fuzzy_char_count=fuzzy_char_count,
                enable_logging=False  # Reduce log noise during search
            )
            
            if result.found:
                if enable_logging:
                    _logger.debug(f"[TextMatch] Found in paragraph {para_idx} using '{result.strategy}' strategy")
                return (para, result)
                
        except Exception as e:
            if enable_logging:
                _logger.debug(f"[TextMatch] Error in paragraph {para_idx}: {e}")
            continue
    
    # Not found in any paragraph
    if enable_logging:
        _logger.debug(f"[TextMatch] Text not found in {len(paragraphs)} paragraphs: '{search_text[:60]}...'")
    
    return (None, TextMatchResult(found=False, strategy='none', search_text=search_text))


class CommentScope(Enum):
    """Scope for which links to comment on."""
    BROKEN_ONLY = "broken_only"          # Only 404, invalid, etc.
    ALL_NON_WORKING = "all_non_working"  # Includes timeout, blocked, unknown


# Status categories
BROKEN_STATUSES = {'broken', 'not_found', '404', 'invalid', 'dns_failed', 'ssl_error'}
NON_WORKING_STATUSES = BROKEN_STATUSES | {'timeout', 'blocked', 'unknown', 'error'}
OK_STATUSES = {'valid', 'working', 'valid_format', 'ok', 'skipped', 'redirect'}


def _truncate_url(url: str, max_len: int = 50) -> str:
    """Truncate URL for logging (privacy/proprietary concerns)."""
    if len(url) <= max_len:
        return url
    return url[:max_len-3] + '...'


def _log_url(url: str, level: str = 'debug'):
    """Log URL with truncation for privacy."""
    truncated = _truncate_url(url)
    getattr(_logger, level)(f"Processing URL: {truncated}")


@dataclass
class HyperlinkLocation:
    """Describes where a hyperlink was found in the document."""
    location_type: str  # 'paragraph', 'table', 'header', 'footer', 'footnote'
    paragraph_index: int = -1
    table_index: int = -1
    row_index: int = -1
    cell_index: int = -1
    section_index: int = -1
    surrounding_text: str = ""  # Context for manual location
    search_string: str = ""     # Exact text to search for
    
    def to_dict(self) -> Dict:
        return {
            'location_type': self.location_type,
            'paragraph_index': self.paragraph_index,
            'table_index': self.table_index,
            'row_index': self.row_index,
            'cell_index': self.cell_index,
            'section_index': self.section_index,
            'surrounding_text': self.surrounding_text,
            'search_string': self.search_string,
        }
    
    def describe(self) -> str:
        """Human-readable location description."""
        if self.location_type == 'table':
            return f"Table {self.table_index + 1}, Row {self.row_index + 1}, Cell {self.cell_index + 1}"
        elif self.location_type == 'header':
            return f"Header (Section {self.section_index + 1})"
        elif self.location_type == 'footer':
            return f"Footer (Section {self.section_index + 1})"
        elif self.location_type == 'footnote':
            return f"Footnote"
        else:
            return f"Paragraph {self.paragraph_index + 1}"


@dataclass
class HyperlinkInfo:
    """Complete information about a hyperlink."""
    target_url: str
    display_text: str
    location: HyperlinkLocation
    element: object = None  # XML element reference
    rel_id: str = ""
    
    def to_dict(self) -> Dict:
        return {
            'target': self.target_url,
            'display_text': self.display_text,
            'location': self.location.to_dict(),
            'rel_id': self.rel_id,
        }


@dataclass
class HyperlinkComment:
    """Represents a comment to be inserted at a hyperlink location."""
    target_url: str
    display_text: str
    status: str
    status_message: str
    location: Optional[HyperlinkLocation] = None
    comment_text: str = ""
    severity: str = "medium"
    
    def __post_init__(self):
        """Generate comment text if not provided."""
        if not self.comment_text:
            self.comment_text = self._generate_comment_text()
    
    def _generate_comment_text(self) -> str:
        """Generate appropriate comment text based on status."""
        status_messages = {
            'broken': f"BROKEN LINK: {self.status_message}\nURL: {self.target_url}",
            'not_found': f"LINK NOT FOUND (404): {self.status_message}\nURL: {self.target_url}",
            '404': f"LINK NOT FOUND (404): Page not found\nURL: {self.target_url}",
            'timeout': f"LINK TIMEOUT: Connection timed out\nURL: {self.target_url}",
            'blocked': f"LINK BLOCKED: Access denied or filtered\nURL: {self.target_url}",
            'dns_failed': f"DNS FAILED: Could not resolve hostname\nURL: {self.target_url}",
            'ssl_error': f"SSL ERROR: Certificate issue\nURL: {self.target_url}",
            'unknown': f"LINK STATUS UNKNOWN: Could not verify\nURL: {self.target_url}",
            'invalid': f"INVALID LINK: {self.status_message}\nURL: {self.target_url}",
            'error': f"LINK ERROR: {self.status_message}\nURL: {self.target_url}",
        }
        
        return status_messages.get(
            self.status.lower(),
            f"LINK ISSUE ({self.status}): {self.status_message}\nURL: {self.target_url}"
        )


def _get_surrounding_text(paragraph, max_chars: int = 60) -> str:
    """Extract surrounding text from a paragraph for context."""
    try:
        text = paragraph.text.strip()
        if len(text) <= max_chars:
            return text
        return text[:max_chars] + "..."
    except Exception:
        return ""


def _find_hyperlinks_in_paragraph(para, para_idx: int, doc_part, location_type: str = 'paragraph',
                                   table_idx: int = -1, row_idx: int = -1, cell_idx: int = -1,
                                   section_idx: int = -1) -> List[HyperlinkInfo]:
    """
    Find hyperlinks in a single paragraph.
    
    Args:
        para: python-docx Paragraph object
        para_idx: Paragraph index in its container
        doc_part: Document part for relationship resolution
        location_type: Where this paragraph lives
        table_idx, row_idx, cell_idx: Table coordinates if applicable
        section_idx: Section index for headers/footers
        
    Returns:
        List of HyperlinkInfo objects
    """
    from docx.oxml.ns import qn
    
    hyperlinks = []
    surrounding_text = _get_surrounding_text(para)
    
    for link_elem in para._element.findall('.//w:hyperlink',
                                            {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
        # Get the relationship ID
        r_id = link_elem.get(qn('r:id'))
        
        # Get display text from runs inside hyperlink
        display_text = ''
        for run in link_elem.findall('.//w:t',
                                     {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
            display_text += run.text or ''
        
        # Resolve URL from relationship
        target_url = ''
        if r_id:
            try:
                rel = doc_part.rels.get(r_id)
                if rel:
                    target_url = rel.target_ref
            except Exception:
                pass
        
        # Create location info
        location = HyperlinkLocation(
            location_type=location_type,
            paragraph_index=para_idx,
            table_index=table_idx,
            row_index=row_idx,
            cell_index=cell_idx,
            section_index=section_idx,
            surrounding_text=surrounding_text,
            search_string=display_text,
        )
        
        hyperlinks.append(HyperlinkInfo(
            target_url=target_url,
            display_text=display_text,
            location=location,
            element=link_elem,
            rel_id=r_id or '',
        ))
    
    return hyperlinks


def find_hyperlinks_in_docx(filepath: str) -> List[HyperlinkInfo]:
    """
    Find ALL hyperlinks in a DOCX document with their locations.
    
    Searches:
    - Body paragraphs
    - Tables (all cells)
    - Headers (all sections)
    - Footers (all sections)
    
    Args:
        filepath: Path to DOCX file
        
    Returns:
        List of HyperlinkInfo objects with full location data
    """
    from docx import Document
    
    doc = Document(filepath)
    all_hyperlinks = []
    
    # 1. Body paragraphs
    _logger.debug("Scanning body paragraphs...")
    for para_idx, para in enumerate(doc.paragraphs):
        links = _find_hyperlinks_in_paragraph(para, para_idx, doc.part, 'paragraph')
        all_hyperlinks.extend(links)
    
    # 2. Tables
    _logger.debug(f"Scanning {len(doc.tables)} tables...")
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                for para_idx, para in enumerate(cell.paragraphs):
                    links = _find_hyperlinks_in_paragraph(
                        para, para_idx, doc.part, 'table',
                        table_idx=table_idx, row_idx=row_idx, cell_idx=cell_idx
                    )
                    all_hyperlinks.extend(links)
    
    # 3. Headers and Footers
    _logger.debug(f"Scanning {len(doc.sections)} section headers/footers...")
    for section_idx, section in enumerate(doc.sections):
        # Header
        try:
            header = section.header
            if header and header.paragraphs:
                for para_idx, para in enumerate(header.paragraphs):
                    links = _find_hyperlinks_in_paragraph(
                        para, para_idx, header.part, 'header',
                        section_idx=section_idx
                    )
                    all_hyperlinks.extend(links)
                # Also check tables in header
                if hasattr(header, 'tables'):
                    for table_idx, table in enumerate(header.tables):
                        for row_idx, row in enumerate(table.rows):
                            for cell_idx, cell in enumerate(row.cells):
                                for para_idx, para in enumerate(cell.paragraphs):
                                    links = _find_hyperlinks_in_paragraph(
                                        para, para_idx, header.part, 'header',
                                        table_idx=table_idx, row_idx=row_idx, 
                                        cell_idx=cell_idx, section_idx=section_idx
                                    )
                                    all_hyperlinks.extend(links)
        except Exception as e:
            _logger.debug(f"Could not scan header for section {section_idx}: {e}")
        
        # Footer
        try:
            footer = section.footer
            if footer and footer.paragraphs:
                for para_idx, para in enumerate(footer.paragraphs):
                    links = _find_hyperlinks_in_paragraph(
                        para, para_idx, footer.part, 'footer',
                        section_idx=section_idx
                    )
                    all_hyperlinks.extend(links)
                # Also check tables in footer
                if hasattr(footer, 'tables'):
                    for table_idx, table in enumerate(footer.tables):
                        for row_idx, row in enumerate(table.rows):
                            for cell_idx, cell in enumerate(row.cells):
                                for para_idx, para in enumerate(cell.paragraphs):
                                    links = _find_hyperlinks_in_paragraph(
                                        para, para_idx, footer.part, 'footer',
                                        table_idx=table_idx, row_idx=row_idx,
                                        cell_idx=cell_idx, section_idx=section_idx
                                    )
                                    all_hyperlinks.extend(links)
        except Exception as e:
            _logger.debug(f"Could not scan footer for section {section_idx}: {e}")
    
    _logger.info(f"Found {len(all_hyperlinks)} hyperlinks in document")
    return all_hyperlinks


def should_comment_link(status: str, scope: CommentScope = CommentScope.ALL_NON_WORKING) -> bool:
    """
    Determine if a link should receive a comment based on status and scope.
    
    Args:
        status: Link status string
        scope: Comment scope setting
        
    Returns:
        True if link should be commented
    """
    status_lower = status.lower()
    
    # Never comment OK links
    if status_lower in OK_STATUSES:
        return False
    
    if scope == CommentScope.BROKEN_ONLY:
        return status_lower in BROKEN_STATUSES
    else:  # ALL_NON_WORKING
        return status_lower in NON_WORKING_STATUSES


def insert_comments_at_hyperlinks(
    filepath: str,
    broken_links: List[Dict],
    author: str = "TechWriterReview",
    output_path: Optional[str] = None,
    scope: CommentScope = CommentScope.ALL_NON_WORKING
) -> Tuple[str, int]:
    """
    Insert comments at broken hyperlink locations in a DOCX file.
    
    Args:
        filepath: Path to source DOCX file
        broken_links: List of broken link info from hyperlink health validation
                      Each dict should have: target, status, status_message, display_text
        author: Comment author name
        output_path: Path for output file (default: adds _commented suffix)
        scope: Which links to comment (broken_only or all_non_working)
        
    Returns:
        Tuple of (output_filepath, comments_inserted_count)
    """
    from docx import Document
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    
    # Open document
    doc = Document(filepath)
    
    # Build lookup of broken links by URL
    broken_by_url = {link.get('target', ''): link for link in broken_links}
    
    # Track inserted comments
    comments_inserted = 0
    comment_id_counter = 0
    
    def process_paragraph(para, doc_part):
        """Process a single paragraph for hyperlinks."""
        nonlocal comments_inserted, comment_id_counter
        
        for link_elem in para._element.findall('.//w:hyperlink',
                                                {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
            # Get relationship ID and resolve URL
            r_id = link_elem.get(qn('r:id'))
            if not r_id:
                continue
            
            try:
                rel = doc_part.rels.get(r_id)
                if not rel:
                    continue
                target_url = rel.target_ref
            except Exception:
                continue
            
            # Check if this link is in our broken list
            if target_url not in broken_by_url:
                continue
            
            broken_info = broken_by_url[target_url]
            status = broken_info.get('status', 'unknown')
            
            # Apply scope filter
            if not should_comment_link(status, scope):
                continue
            
            # Create comment
            comment_obj = HyperlinkComment(
                target_url=target_url,
                display_text=broken_info.get('display_text', ''),
                status=status,
                status_message=broken_info.get('status_message', 'Link issue detected')
            )
            
            # Add comment marks to element
            try:
                _add_comment_marks_to_element(link_elem, comment_id_counter, comment_obj.comment_text, author)
                comments_inserted += 1
                comment_id_counter += 1
                _log_url(target_url, 'debug')
            except Exception as e:
                _logger.warning(f"Could not add comment: {e}")
    
    # Process body paragraphs
    for para in doc.paragraphs:
        process_paragraph(para, doc.part)
    
    # Process tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    process_paragraph(para, doc.part)
    
    # Process headers and footers
    for section in doc.sections:
        try:
            if section.header and section.header.paragraphs:
                for para in section.header.paragraphs:
                    process_paragraph(para, section.header.part)
                if hasattr(section.header, 'tables'):
                    for table in section.header.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for para in cell.paragraphs:
                                    process_paragraph(para, section.header.part)
        except Exception:
            pass
        
        try:
            if section.footer and section.footer.paragraphs:
                for para in section.footer.paragraphs:
                    process_paragraph(para, section.footer.part)
                if hasattr(section.footer, 'tables'):
                    for table in section.footer.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for para in cell.paragraphs:
                                    process_paragraph(para, section.footer.part)
        except Exception:
            pass
    
    # Determine output path
    if not output_path:
        path = Path(filepath)
        output_path = str(path.parent / f"{path.stem}_commented{path.suffix}")
    
    # Save document
    doc.save(output_path)
    
    _logger.info(f"Inserted {comments_inserted} comments, saved to {_truncate_url(output_path, 80)}")
    return output_path, comments_inserted


def _add_comment_marks_to_element(element, comment_id: int, comment_text: str, author: str):
    """
    Add comment range start/end marks around an element's content.
    
    Note: This adds the XML markers for comments. Full comment insertion
    requires also adding to the comments.xml part, which python-docx
    doesn't fully support. The markers allow Word to show comment indicators.
    
    Args:
        element: The XML element to wrap with comment marks
        comment_id: The comment ID to use
        comment_text: Text content of the comment (stored as attribute for potential use)
        author: Comment author name
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    
    # Create commentRangeStart element
    comment_start = OxmlElement('w:commentRangeStart')
    comment_start.set(qn('w:id'), str(comment_id))
    
    # Create commentRangeEnd element
    comment_end = OxmlElement('w:commentRangeEnd')
    comment_end.set(qn('w:id'), str(comment_id))
    
    # Create commentReference element (goes in a run)
    comment_ref_run = OxmlElement('w:r')
    comment_ref = OxmlElement('w:commentReference')
    comment_ref.set(qn('w:id'), str(comment_id))
    comment_ref_run.append(comment_ref)
    
    # Insert start before the element's first child
    if len(element) > 0:
        element.insert(0, comment_start)
    else:
        element.append(comment_start)
    
    # Append end and reference after the element's content
    element.append(comment_end)
    element.append(comment_ref_run)


def generate_comment_pack(
    broken_links: List[Dict],
    document_name: str = "Document",
    hyperlink_info: Optional[List[HyperlinkInfo]] = None,
    scope: CommentScope = CommentScope.ALL_NON_WORKING
) -> str:
    """
    Generate a text file listing all broken links with suggested comments
    and location hints for manual application.
    
    Args:
        broken_links: List of broken link info from hyperlink health validation
        document_name: Name of the source document
        hyperlink_info: Optional list of HyperlinkInfo with location data
        scope: Which links to include
        
    Returns:
        Text content of the comment pack
    """
    # Build location lookup if provided
    location_by_url = {}
    if hyperlink_info:
        for info in hyperlink_info:
            location_by_url[info.target_url] = info.location
    
    # Count commentable links
    commentable_count = len([l for l in broken_links if should_comment_link(l.get('status', ''), scope)])
    
    lines = [
        "=" * 70,
        "HYPERLINK COMMENT PACK",
        f"Document: {document_name}",
        f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        f"Comment Scope: {scope.value}",
        f"Total Issues: {commentable_count}",
        "=" * 70,
        "",
        "Instructions:",
        "1. Open your document in Microsoft Word",
        "2. Use the LOCATION hints below to find each hyperlink",
        "3. Select the hyperlink text (shown in SEARCH FOR)",
        "4. Insert > Comment (or Ctrl+Alt+M)",
        "5. Copy the suggested comment text",
        "",
        "-" * 70,
        ""
    ]
    
    entry_num = 0
    for link in broken_links:
        target = link.get('target', 'Unknown URL')
        display = link.get('display_text', 'Unknown')
        status = link.get('status', 'unknown')
        message = link.get('status_message', 'Issue detected')
        
        # Apply scope filter
        if not should_comment_link(status, scope):
            continue
        
        entry_num += 1
        
        comment_obj = HyperlinkComment(
            target_url=target,
            display_text=display,
            status=status,
            status_message=message
        )
        
        # Get location info if available
        # v3.0.105: BUG-005 FIX - Also check for hyperlink_info inside the link dict
        location = location_by_url.get(target)
        if not location and 'hyperlink_info' in link:
            # hyperlink_info might be a dict with 'location' key
            hi = link.get('hyperlink_info', {})
            if isinstance(hi, dict) and 'location' in hi:
                location = hi['location']
            elif hasattr(hi, 'location'):
                location = hi.location
        
        lines.extend([
            f"[{entry_num}] STATUS: {status.upper()}",
            f"    SEARCH FOR: \"{display}\"",
            f"    URL: {target[:80]}{'...' if len(target) > 80 else ''}",
        ])
        
        # Add location hints
        if location:
            lines.append(f"    LOCATION: {location.describe()}")
            # v3.0.105: BUG-005 FIX - Get surrounding_text from location or hyperlink_info
            surrounding_text = getattr(location, 'surrounding_text', None)
            if not surrounding_text and 'hyperlink_info' in link:
                hi = link.get('hyperlink_info', {})
                if isinstance(hi, dict):
                    surrounding_text = hi.get('surrounding_text')
            if surrounding_text:
                ctx = surrounding_text[:60]
                lines.append(f"    CONTEXT: \"{ctx}{'...' if len(surrounding_text) > 60 else ''}\"")
        else:
            lines.append(f"    LOCATION: Use Ctrl+F to search for the link text")
        
        lines.extend([
            f"    ",
            f"    SUGGESTED COMMENT:",
            f"    {'-' * 40}",
        ])
        
        # Indent comment text
        for line in comment_obj.comment_text.split('\n'):
            lines.append(f"    {line}")
        
        lines.extend([
            f"    {'-' * 40}",
            ""
        ])
    
    if entry_num == 0:
        lines.extend([
            "No links match the selected comment scope.",
            f"Scope: {scope.value}",
            ""
        ])
    
    lines.extend([
        "=" * 70,
        "END OF COMMENT PACK",
        "=" * 70
    ])
    
    return '\n'.join(lines)


def export_comment_pack(
    broken_links: List[Dict],
    output_path: str,
    document_name: str = "Document",
    hyperlink_info: Optional[List[HyperlinkInfo]] = None,
    scope: CommentScope = CommentScope.ALL_NON_WORKING
) -> str:
    """
    Export comment pack to a text file.
    
    Args:
        broken_links: List of broken link info
        output_path: Path for output text file
        document_name: Name of source document
        hyperlink_info: Optional location info
        scope: Comment scope
        
    Returns:
        Path to created file
    """
    content = generate_comment_pack(broken_links, document_name, hyperlink_info, scope)
    
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(content)
    
    return output_path


def process_hyperlink_health_results(
    docx_path: str,
    health_report: Dict,
    mode: str = 'insert',
    author: str = 'TechWriterReview',
    output_dir: Optional[str] = None,
    scope: str = 'all_non_working'  # 'broken_only' or 'all_non_working'
) -> Dict:
    """
    Process hyperlink health results and either insert comments or generate pack.
    
    Args:
        docx_path: Path to source DOCX file
        health_report: Hyperlink health report dict from validation
        mode: 'insert' for direct DOCX comments, 'pack' for text file
        author: Comment author name
        output_dir: Directory for output files (default: same as source)
        scope: 'broken_only' or 'all_non_working' (default)
        
    Returns:
        Result dict with output paths and counts
    """
    # Parse scope
    comment_scope = CommentScope.BROKEN_ONLY if scope == 'broken_only' else CommentScope.ALL_NON_WORKING
    
    # Extract links from report
    links = health_report.get('links', [])
    
    # Filter to commentable links
    commentable_links = [
        link for link in links
        if should_comment_link(link.get('status', ''), comment_scope)
    ]
    
    if not commentable_links:
        return {
            'success': True,
            'mode': mode,
            'scope': scope,
            'message': 'No links match the comment scope - no comments needed',
            'total_links': len(links),
            'commentable_count': 0,
            'output_path': None
        }
    
    # Get full hyperlink info with locations for better comment pack
    hyperlink_info = None
    try:
        hyperlink_info = find_hyperlinks_in_docx(docx_path)
    except Exception as e:
        _logger.warning(f"Could not extract hyperlink locations: {e}")
    
    # Determine output directory
    if not output_dir:
        output_dir = str(Path(docx_path).parent)
    
    doc_name = Path(docx_path).stem
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    
    if mode == 'insert':
        # Direct DOCX comment insertion
        output_path = os.path.join(output_dir, f"{doc_name}_commented_{timestamp}.docx")
        try:
            result_path, count = insert_comments_at_hyperlinks(
                filepath=docx_path,
                broken_links=links,  # Pass all links, filtering happens inside
                author=author,
                output_path=output_path,
                scope=comment_scope
            )
            return {
                'success': True,
                'mode': 'insert',
                'scope': scope,
                'message': f'Inserted {count} comments into document',
                'total_links': len(links),
                'commentable_count': len(commentable_links),
                'comments_inserted': count,
                'output_path': result_path
            }
        except Exception as e:
            _logger.error(f"Comment insertion failed: {e}")
            return {
                'success': False,
                'mode': 'insert',
                'scope': scope,
                'error': str(e),
                'total_links': len(links),
                'commentable_count': len(commentable_links),
                'output_path': None
            }
    
    else:  # mode == 'pack'
        # Generate comment pack text file
        output_path = os.path.join(output_dir, f"{doc_name}_comment_pack_{timestamp}.txt")
        try:
            result_path = export_comment_pack(
                broken_links=links,  # Pass all, filtering in generate
                output_path=output_path,
                document_name=Path(docx_path).name,
                hyperlink_info=hyperlink_info,
                scope=comment_scope
            )
            return {
                'success': True,
                'mode': 'pack',
                'scope': scope,
                'message': f'Generated comment pack with {len(commentable_links)} entries',
                'total_links': len(links),
                'commentable_count': len(commentable_links),
                'output_path': result_path
            }
        except Exception as e:
            _logger.error(f"Comment pack generation failed: {e}")
            return {
                'success': False,
                'mode': 'pack',
                'scope': scope,
                'error': str(e),
                'total_links': len(links),
                'commentable_count': len(commentable_links),
                'output_path': None
            }


# =============================================================================
# MODULE TEST
# =============================================================================

if __name__ == "__main__":
    print(f"Comment Inserter Module v{__version__}")
    print("=" * 60)
    
    # v3.0.109: Test text normalization functions
    print("\n--- v3.0.109: Text Normalization Tests ---")
    
    # Test quote normalization
    print("\n1. Quote Normalization:")
    test_quotes = '\u201cHello\u201d said the \u201cdeveloper\u201d'  # Smart quotes (U+201C, U+201D)
    normalized = normalize_quotes(test_quotes)
    print(f"   Input:  {repr(test_quotes)}")
    print(f"   Output: {repr(normalized)}")
    print(f"   ✓ Pass" if '\u201c' not in normalized and '\u201d' not in normalized else "   ✗ Fail")
    
    # Test single quotes
    test_single = '\u2018It\u2019s working\u2019'  # Smart single quotes (U+2018, U+2019)
    norm_single = normalize_quotes(test_single)
    print(f"   Input:  {repr(test_single)}")
    print(f"   Output: {repr(norm_single)}")
    print(f"   ✓ Pass" if '\u2018' not in norm_single and '\u2019' not in norm_single else "   ✗ Fail")
    
    # Test whitespace normalization
    print("\n2. Whitespace Normalization:")
    test_ws = "Multiple   spaces\t\ttabs\nand\nnewlines"
    norm_ws = normalize_whitespace(test_ws)
    print(f"   Input:  {repr(test_ws)}")
    print(f"   Output: {repr(norm_ws)}")
    expected_ws = "Multiple spaces tabs and newlines"
    print(f"   ✓ Pass" if norm_ws == expected_ws else f"   ✗ Fail (expected: {repr(expected_ws)})")
    
    # Test full normalization
    print("\n3. Full Text Normalization:")
    test_full = '\u201cSmart quotes\u201d   with    weird  spaces'  # Smart quotes + multiple spaces
    norm_full = normalize_text_for_matching(test_full)
    print(f"   Input:  {repr(test_full)}")
    print(f"   Output: {repr(norm_full)}")
    print(f"   ✓ Pass" if '\u201c' not in norm_full and '  ' not in norm_full else "   ✗ Fail")
    
    # v3.0.109: Test text matching strategies
    print("\n--- v3.0.109: Text Matching Strategy Tests ---")
    
    # Test exact match
    print("\n4. Exact Match:")
    doc_text = "This document contains exact matching text here."
    result = find_text_in_document("exact matching text", doc_text, enable_logging=False)
    print(f"   Search: 'exact matching text' in document")
    print(f"   Result: {result}")
    print(f"   ✓ Pass" if result.found and result.strategy == 'exact' else "   ✗ Fail")
    
    # Test normalized match (quote mismatch)
    print("\n5. Normalized Match (quote mismatch):")
    doc_text_quotes = 'The document says \u201chello world\u201d with smart quotes.'  # Smart quotes in doc
    search_straight = 'The document says "hello world" with smart quotes.'  # Straight quotes in search
    result_norm = find_text_in_document(search_straight, doc_text_quotes, enable_logging=False)
    print(f"   Doc has smart quotes, search has straight quotes")
    print(f"   Result: {result_norm}")
    print(f"   ✓ Pass" if result_norm.found and result_norm.strategy == 'normalized' else "   ✗ Fail")
    
    # Test normalized match (whitespace mismatch)
    print("\n6. Normalized Match (whitespace mismatch):")
    doc_text_ws = "Text with   multiple    spaces between words."
    search_ws = "Text with multiple spaces between words."
    result_ws = find_text_in_document(search_ws, doc_text_ws, enable_logging=False)
    print(f"   Doc has multiple spaces, search has single spaces")
    print(f"   Result: {result_ws}")
    print(f"   ✓ Pass" if result_ws.found and result_ws.strategy == 'normalized' else "   ✗ Fail")
    
    # Test fuzzy match
    print("\n7. Fuzzy Match (truncated text):")
    doc_text_fuzzy = "This is a very long sentence that continues for quite a while and has more content."
    search_fuzzy = "This is a very long sentence that continues differently in the search text"
    result_fuzzy = find_text_in_document(search_fuzzy, doc_text_fuzzy, fuzzy_min_length=20, fuzzy_char_count=30, enable_logging=False)
    print(f"   Search first 30 chars of long text")
    print(f"   Result: {result_fuzzy}")
    print(f"   ✓ Pass" if result_fuzzy.found and result_fuzzy.strategy == 'fuzzy' else "   ✗ Fail")
    
    # Test no match
    print("\n8. No Match:")
    doc_text_no = "This document has completely different content."
    result_no = find_text_in_document("nonexistent text here", doc_text_no, enable_logging=False)
    print(f"   Search for text that doesn't exist")
    print(f"   Result: {result_no}")
    print(f"   ✓ Pass" if not result_no.found and result_no.strategy == 'none' else "   ✗ Fail")
    
    # Test TextMatchResult boolean behavior
    print("\n9. TextMatchResult Boolean:")
    result_true = TextMatchResult(found=True, strategy='exact')
    result_false = TextMatchResult(found=False, strategy='none')
    print(f"   bool(found=True):  {bool(result_true)}")
    print(f"   bool(found=False): {bool(result_false)}")
    print(f"   ✓ Pass" if result_true and not result_false else "   ✗ Fail")
    
    print("\n" + "=" * 60)
    
    # Original tests
    print("\n--- Original Module Tests ---")
    
    # Test comment generation
    comment = HyperlinkComment(
        target_url="https://example.com/broken",
        display_text="Example Link",
        status="broken",
        status_message="HTTP 404"
    )
    print(f"\nSample comment:\n{comment.comment_text}")
    
    # Test scope filtering
    print("\nScope filtering tests:")
    print(f"  broken + BROKEN_ONLY: {should_comment_link('broken', CommentScope.BROKEN_ONLY)}")  # True
    print(f"  timeout + BROKEN_ONLY: {should_comment_link('timeout', CommentScope.BROKEN_ONLY)}")  # False
    print(f"  timeout + ALL_NON_WORKING: {should_comment_link('timeout', CommentScope.ALL_NON_WORKING)}")  # True
    print(f"  valid + ALL_NON_WORKING: {should_comment_link('valid', CommentScope.ALL_NON_WORKING)}")  # False
    
    # Test location description
    loc = HyperlinkLocation(
        location_type='table',
        table_index=0,
        row_index=2,
        cell_index=1,
        surrounding_text="This is the context text around the link"
    )
    print(f"\nLocation description: {loc.describe()}")
    
    print("\n" + "=" * 60)
    print("✓ Comment Inserter module ready (v3.0.109 with improved text matching)")
