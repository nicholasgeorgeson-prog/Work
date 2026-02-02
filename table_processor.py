"""
TechWriterReview - Table-Aware Document Processor
=================================================

Lightweight alternative to Docling for extracting roles from documents with tables.

Dependencies (all lightweight):
- pdfplumber: Table extraction from PDFs (~5 MB)
- python-docx: Word document handling (already in TWR)
- spaCy (optional): NLP for role extraction (~82 MB)

Total: ~87 MB vs Docling's ~1.5 GB

Usage:
    from table_processor import extract_tables_and_roles
    
    result = extract_tables_and_roles("document.pdf")
    print(result['tables'])  # Structured table data
    print(result['roles'])   # Extracted roles
"""

import re
from typing import Dict, List, Optional, Any
from pathlib import Path

# Try imports
try:
    import pdfplumber
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False
    print("[TableProcessor] pdfplumber not installed - PDF support disabled")

try:
    from docx import Document as DocxDocument
    from docx.table import Table as DocxTable
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False
    print("[TableProcessor] python-docx not installed - DOCX support disabled")


class TableProcessor:
    """
    Extracts tables from PDFs and Word documents, then identifies roles.
    
    Much lighter than Docling (~87 MB vs ~1.5 GB) while handling
    the most common table formats in work instructions.
    """
    
    def __init__(self):
        # Common table headers that indicate role columns
        self.ROLE_HEADERS = {
            'role', 'roles', 'responsibility', 'responsibilities',
            'responsible', 'accountable', 'consulted', 'informed',
            'performed by', 'approved by', 'reviewed by', 'prepared by',
            'owner', 'assignee', 'assigned to', 'personnel', 'person',
            'who', 'actor', 'function', 'position', 'title', 'job title',
            'r', 'a', 'c', 'i',  # RACI columns
        }
        
        # Headers that indicate a RACI matrix
        self.RACI_HEADERS = {'r', 'a', 'c', 'i', 'responsible', 'accountable', 'consulted', 'informed'}
    
    def process_document(self, filepath: str) -> Dict[str, Any]:
        """
        Process a document and extract tables with roles.
        
        Args:
            filepath: Path to PDF or DOCX file
            
        Returns:
            Dict with 'tables', 'roles', 'text', 'stats'
        """
        path = Path(filepath)
        ext = path.suffix.lower()
        
        if ext == '.pdf':
            if not PDF_SUPPORT:
                return {'error': 'pdfplumber not installed'}
            return self._process_pdf(filepath)
        elif ext in ('.docx', '.doc'):
            if not DOCX_SUPPORT:
                return {'error': 'python-docx not installed'}
            return self._process_docx(filepath)
        else:
            return {'error': f'Unsupported file type: {ext}'}
    
    def _process_pdf(self, filepath: str) -> Dict[str, Any]:
        """Extract tables and text from PDF with optimized settings."""
        tables = []
        all_text = []
        roles_from_tables = set()
        
        # Optimized settings for text/table extraction only
        # No image extraction, no curve detection (faster parsing)
        laparams = {
            'line_margin': 0.5,
            'word_margin': 0.1,
            'char_margin': 2.0,
            'boxes_flow': 0.5,
        }
        
        with pdfplumber.open(filepath) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                # Extract tables with optimized settings
                # Disable image extraction completely
                table_settings = {
                    "vertical_strategy": "lines",
                    "horizontal_strategy": "lines",
                    "snap_tolerance": 3,
                    "join_tolerance": 3,
                    "edge_min_length": 10,
                    "min_words_vertical": 3,
                    "min_words_horizontal": 1,
                }
                
                try:
                    page_tables = page.extract_tables(table_settings)
                except Exception:
                    # Fallback to default if custom settings fail
                    page_tables = page.extract_tables()
                
                for table_idx, table in enumerate(page_tables or []):
                    if table and len(table) > 1:  # Has header + data
                        processed = self._process_table(table, page_num, table_idx)
                        if processed:
                            tables.append(processed)
                            roles_from_tables.update(processed.get('roles', []))
                
                # Extract text only (no images) - filter out image objects
                # Use filtered page to avoid image processing overhead
                text = page.extract_text(layout=False) or ''
                all_text.append(text)
        
        return {
            'success': True,
            'tables': tables,
            'roles_from_tables': list(roles_from_tables),
            'text': '\n'.join(all_text),
            'stats': {
                'pages': len(all_text),
                'tables_found': len(tables),
                'roles_in_tables': len(roles_from_tables),
            }
        }
    
    def _process_docx(self, filepath: str) -> Dict[str, Any]:
        """Extract tables and text from Word document."""
        doc = DocxDocument(filepath)
        tables = []
        all_text = []
        roles_from_tables = set()
        
        # Extract tables
        for table_idx, table in enumerate(doc.tables):
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            
            if table_data and len(table_data) > 1:
                processed = self._process_table(table_data, 1, table_idx)
                if processed:
                    tables.append(processed)
                    roles_from_tables.update(processed.get('roles', []))
        
        # Extract text from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                all_text.append(para.text)
        
        return {
            'success': True,
            'tables': tables,
            'roles_from_tables': list(roles_from_tables),
            'text': '\n'.join(all_text),
            'stats': {
                'tables_found': len(tables),
                'roles_in_tables': len(roles_from_tables),
            }
        }
    
    def _process_table(self, table: List[List[str]], page: int, idx: int) -> Optional[Dict]:
        """
        Process a single table and extract roles.
        
        Handles:
        - RACI matrices
        - Role assignment tables
        - Responsibility tables
        """
        if not table or len(table) < 2:
            return None
        
        # Clean headers
        headers = [str(h).strip().lower() if h else '' for h in table[0]]
        
        # Detect table type
        is_raci = self._is_raci_table(headers)
        role_columns = self._find_role_columns(headers)
        
        roles = set()
        structured_data = []
        
        # Process data rows
        for row_idx, row in enumerate(table[1:], 1):
            row_data = {}
            
            for col_idx, cell in enumerate(row):
                if col_idx < len(headers):
                    header = headers[col_idx]
                    value = str(cell).strip() if cell else ''
                    row_data[header or f'col_{col_idx}'] = value
                    
                    # Extract roles from role columns
                    if col_idx in role_columns and value:
                        # Split on common delimiters
                        cell_roles = re.split(r'[,;/\n]|\band\b|\bor\b', value)
                        for role in cell_roles:
                            role = role.strip()
                            if role and len(role) > 2 and not role.isdigit():
                                if self._looks_like_role(role):
                                    roles.add(role)
                    
                    # For RACI tables, extract roles from row headers
                    if is_raci and col_idx == 0 and value:
                        # First column usually has activity, but check for roles
                        if self._looks_like_role(value):
                            roles.add(value)
            
            if row_data:
                structured_data.append(row_data)
        
        # For RACI tables, headers often ARE the roles
        if is_raci:
            for header in headers:
                if header and self._looks_like_role(header) and header not in self.RACI_HEADERS:
                    roles.add(header.title())
        
        return {
            'page': page,
            'index': idx,
            'type': 'RACI' if is_raci else 'general',
            'headers': headers,
            'data': structured_data,
            'roles': list(roles),
            'row_count': len(structured_data),
        }
    
    def _is_raci_table(self, headers: List[str]) -> bool:
        """Check if table is a RACI matrix."""
        header_set = set(h.strip().lower() for h in headers if h)
        raci_matches = header_set & self.RACI_HEADERS
        return len(raci_matches) >= 2
    
    def _find_role_columns(self, headers: List[str]) -> List[int]:
        """Find column indices that contain roles."""
        role_cols = []
        for idx, header in enumerate(headers):
            if header and header.lower() in self.ROLE_HEADERS:
                role_cols.append(idx)
        return role_cols
    
    def _looks_like_role(self, text: str) -> bool:
        """Check if text looks like a role name."""
        text = text.strip()
        
        # Too short or too long
        if len(text) < 3 or len(text) > 50:
            return False
        
        # Is a number or date
        if text.isdigit() or re.match(r'^\d{1,2}[/-]\d{1,2}', text):
            return False
        
        # Single words that are likely tasks, not roles
        task_words = {
            'assembly', 'inspection', 'testing', 'review', 'analysis',
            'design', 'development', 'integration', 'verification', 'validation',
            'procurement', 'shipping', 'receiving', 'planning', 'scheduling',
            'manufacturing', 'production', 'maintenance', 'installation',
        }
        if text.lower() in task_words:
            return False
        
        # Common role suffixes (strong indicator)
        role_indicators = [
            'engineer', 'manager', 'director', 'supervisor', 'lead',
            'technician', 'analyst', 'specialist', 'coordinator', 'officer',
            'inspector', 'auditor', 'reviewer', 'owner', 'authority',
            'clerk', 'assistant', 'administrator', 'planner', 'scientist',
        ]
        text_lower = text.lower()
        if any(ind in text_lower for ind in role_indicators):
            return True
        
        # Title case with 2-4 words (more likely a role than single word)
        words = text.split()
        if 2 <= len(words) <= 4 and text[0].isupper():
            return True
        
        # Single title-case word - need role indicator
        if len(words) == 1:
            return False
        
        return False


def extract_tables_and_roles(filepath: str) -> Dict[str, Any]:
    """
    Convenience function to extract tables and roles from a document.
    
    Args:
        filepath: Path to PDF or DOCX
        
    Returns:
        Dict with extracted tables, roles, and full text
    """
    processor = TableProcessor()
    return processor.process_document(filepath)


# =============================================================================
# Integration with Role Extractor V3
# =============================================================================

def extract_roles_with_tables(filepath: str, use_spacy: bool = True) -> Dict[str, Any]:
    """
    Full extraction pipeline: tables + NLP.
    
    1. Extracts tables and identifies roles in table cells
    2. Extracts full text 
    3. Runs NLP-based role extraction on text
    4. Combines results
    
    Args:
        filepath: Path to document
        use_spacy: Whether to use spaCy for text analysis
        
    Returns:
        Combined results from table and text extraction
    """
    # Import role extractor
    try:
        from role_extractor_v3 import extract_roles_from_text
        NLP_AVAILABLE = True
    except ImportError:
        NLP_AVAILABLE = False
    
    # Process document for tables
    processor = TableProcessor()
    doc_result = processor.process_document(filepath)
    
    if not doc_result.get('success'):
        return doc_result
    
    # Get roles from tables
    table_roles = set(doc_result.get('roles_from_tables', []))
    
    # Get roles from text using NLP
    text_roles = {}
    if NLP_AVAILABLE and doc_result.get('text'):
        nlp_result = extract_roles_from_text(doc_result['text'], use_spacy=use_spacy)
        text_roles = nlp_result.get('roles', {})
    
    # Combine results
    all_roles = {}
    
    # Add table roles with high confidence (explicit in document)
    for role in table_roles:
        all_roles[role] = {
            'canonical_name': role,
            'confidence': 0.95,
            'source': 'table',
            'detection_methods': ['table_extraction'],
        }
    
    # Add text roles (may overlap with table roles)
    for role, data in text_roles.items():
        if role in all_roles:
            # Boost confidence if found in both
            all_roles[role]['confidence'] = min(0.99, all_roles[role]['confidence'] + 0.05)
            all_roles[role]['detection_methods'].extend(data.get('detection_methods', []))
            all_roles[role]['source'] = 'both'
        else:
            all_roles[role] = {
                'canonical_name': data.get('canonical_name', role),
                'confidence': data.get('confidence', 0.7),
                'source': 'text',
                'detection_methods': data.get('detection_methods', []),
            }
    
    return {
        'success': True,
        'roles': all_roles,
        'tables': doc_result.get('tables', []),
        'stats': {
            'tables_found': len(doc_result.get('tables', [])),
            'roles_from_tables': len(table_roles),
            'roles_from_text': len(text_roles),
            'total_unique_roles': len(all_roles),
        }
    }


# =============================================================================
# TEST
# =============================================================================

if __name__ == "__main__":
    print("=" * 60)
    print("TABLE-AWARE DOCUMENT PROCESSOR")
    print("Lightweight alternative to Docling (~87 MB vs ~1.5 GB)")
    print("=" * 60)
    
    # Test with sample table
    sample_table = [
        ['Task', 'Responsible', 'Accountable', 'Consulted', 'Informed'],
        ['Assembly', 'Manufacturing Technician', 'Production Supervisor', 'Quality Engineer', 'Project Manager'],
        ['Inspection', 'Quality Inspector', 'QA Manager', 'Manufacturing Engineer', 'Test Lead'],
        ['Testing', 'Test Technician', 'Test Lead', 'Systems Engineer', 'Quality Assurance'],
    ]
    
    processor = TableProcessor()
    result = processor._process_table(sample_table, 1, 0)
    
    print(f"\nSample RACI Table Analysis:")
    print(f"  Type: {result['type']}")
    print(f"  Roles found: {len(result['roles'])}")
    print(f"\n  Extracted roles:")
    for role in sorted(result['roles']):
        print(f"    - {role}")
