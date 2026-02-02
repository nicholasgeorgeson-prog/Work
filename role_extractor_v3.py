"""
Role Extractor for Engineering Work Instructions
Version 3.0 - Production-Ready for Air-Gapped Networks

Extracts organizational roles and their associated responsibilities from 
engineering documents (Word docs, PDFs, text files) using pattern matching
and linguistic rules. No external AI/API dependencies.

Dependencies (commonly available on closed networks):
- re (standard library)
- collections (standard library)  
- dataclasses (standard library, Python 3.7+)
- python-docx (pip install python-docx) - for Word documents
- PyPDF2 or pdfplumber (pip install PyPDF2 pdfplumber) - for PDFs

Author: Nick / SAIC Systems Engineering
For use with: Technical Review Tool
"""

import re
from collections import defaultdict
from dataclasses import dataclass, field
from typing import List, Dict, Set, Tuple, Optional
import os
import csv
from enum import Enum

# Structured logging support
try:
    from config_logging import get_logger
    _logger = get_logger('role_extractor')
except ImportError:
    _logger = None

def _log(message: str, level: str = 'info', **kwargs):
    """Internal logging helper with fallback."""
    if _logger:
        getattr(_logger, level)(message, **kwargs)
    elif level in ('warning', 'error', 'critical'):
        print(f"[RoleExtractor] {level.upper()}: {message}")


# v3.0.100: ReDoS Protection (ISSUE-001)
# Maximum input length for regex operations to prevent CPU exhaustion
MAX_REGEX_INPUT_LENGTH = 10000
REGEX_CHUNK_SIZE = 5000  # Process in chunks if needed


def safe_regex_search(pattern, text: str, flags=0, max_length: int = MAX_REGEX_INPUT_LENGTH):
    """
    Safe wrapper for regex search with input length limiting.
    Prevents ReDoS attacks by truncating overly long inputs.
    
    Args:
        pattern: Compiled regex or pattern string
        text: Text to search
        flags: Regex flags (if pattern is string)
        max_length: Maximum input length to process
        
    Returns:
        Match object or None
    """
    if not text:
        return None
    
    # Truncate if too long
    if len(text) > max_length:
        _log(f"Input truncated from {len(text)} to {max_length} chars for regex safety", level='debug')
        text = text[:max_length]
    
    try:
        if isinstance(pattern, str):
            return re.search(pattern, text, flags)
        return pattern.search(text)
    except re.error as e:
        _log(f"Regex error: {e}", level='warning')
        return None


def safe_regex_findall(pattern, text: str, flags=0, max_length: int = MAX_REGEX_INPUT_LENGTH):
    """
    Safe wrapper for regex findall with input length limiting.
    
    Args:
        pattern: Compiled regex or pattern string
        text: Text to search
        flags: Regex flags (if pattern is string)
        max_length: Maximum input length to process
        
    Returns:
        List of matches (empty if error)
    """
    if not text:
        return []
    
    # Truncate if too long
    if len(text) > max_length:
        _log(f"Input truncated from {len(text)} to {max_length} chars for regex safety", level='debug')
        text = text[:max_length]
    
    try:
        if isinstance(pattern, str):
            return re.findall(pattern, text, flags)
        return pattern.findall(text)
    except re.error as e:
        _log(f"Regex error: {e}", level='warning')
        return []


def safe_regex_finditer(pattern, text: str, flags=0, max_length: int = MAX_REGEX_INPUT_LENGTH):
    """
    Safe wrapper for regex finditer with input length limiting.
    
    Args:
        pattern: Compiled regex or pattern string
        text: Text to search
        flags: Regex flags (if pattern is string)
        max_length: Maximum input length to process
        
    Yields:
        Match objects
    """
    if not text:
        return
    
    # Truncate if too long
    if len(text) > max_length:
        _log(f"Input truncated from {len(text)} to {max_length} chars for regex safety", level='debug')
        text = text[:max_length]
    
    try:
        if isinstance(pattern, str):
            yield from re.finditer(pattern, text, flags)
        else:
            yield from pattern.finditer(text)
    except re.error as e:
        _log(f"Regex error: {e}", level='warning')


class EntityKind(Enum):
    """Classification of extracted entity type."""
    ROLE = "role"
    DELIVERABLE = "deliverable"
    UNKNOWN = "unknown"


@dataclass
class RoleOccurrence:
    """Represents a single occurrence of a role in the document."""
    role: str
    context: str
    responsibility: str
    action_type: str
    location: str
    confidence: float


@dataclass 
class ExtractedRole:
    """Aggregated information about a discovered entity (role or deliverable)."""
    canonical_name: str
    entity_kind: EntityKind = EntityKind.UNKNOWN  # v3.0.12: Added entity classification
    kind_confidence: float = 0.0  # v3.0.12: Confidence in classification
    kind_reason: str = ""  # v3.0.12: Why classified this way
    variants: Set[str] = field(default_factory=set)
    occurrences: List[RoleOccurrence] = field(default_factory=list)
    responsibilities: List[str] = field(default_factory=list)
    action_types: Dict[str, int] = field(default_factory=lambda: defaultdict(int))
    
    @property
    def frequency(self) -> int:
        return len(self.occurrences)
    
    @property
    def avg_confidence(self) -> float:
        if not self.occurrences:
            return 0.0
        return sum(o.confidence for o in self.occurrences) / len(self.occurrences)
    
    def to_dict(self) -> dict:
        """Convert to dictionary for JSON serialization."""
        return {
            'canonical_name': self.canonical_name,
            'entity_kind': self.entity_kind.value,
            'kind_confidence': self.kind_confidence,
            'kind_reason': self.kind_reason,
            'variants': list(self.variants),
            'frequency': self.frequency,
            'avg_confidence': self.avg_confidence,
            'responsibilities': self.responsibilities,
            'action_types': dict(self.action_types)
        }


class RoleExtractor:
    """
    Extracts organizational roles from engineering documents.
    Pure Python implementation for air-gapped networks.
    """
    
    # =========================================================================
    # CONFIGURATION - Customize these for your organization
    # =========================================================================
    
    ROLE_SUFFIXES = [
        'engineer', 'manager', 'lead', 'director', 'officer', 'specialist',
        'analyst', 'coordinator', 'administrator', 'authority', 'chief',
        'supervisor', 'inspector', 'auditor', 'reviewer', 'approver',
        'representative', 'owner', 'custodian', 'architect', 'integrator',
        'technician', 'scientist', 'investigator', 'controller', 'planner',
        'panel', 'board', 'council', 'committee', 'team', 'group'
    ]
    
    ROLE_MODIFIERS = [
        'project', 'program', 'systems', 'system', 'lead', 'chief', 'senior',
        'deputy', 'assistant', 'associate', 'principal', 'technical', 'quality',
        'safety', 'mission', 'flight', 'ground', 'test', 'integration',
        'verification', 'validation', 'configuration', 'data', 'risk',
        'requirements', 'interface', 'software', 'hardware', 'mechanical',
        'electrical', 'structural', 'thermal', 'propulsion', 'avionics',
        'reliability', 'maintainability', 'logistics', 'operations', 'security',
        'environmental', 'human', 'factors', 'design', 'manufacturing', 'production',
        'subsystem', 'component', 'element', 'responsible', 'cognizant',
        'designated', 'authorized', 'certifying', 'contracting', 'procurement',
        'review', 'control', 'change', 'engineering', 'standing', 'independent',
        'working', 'action', 'steering', 'executive', 'advisory'
    ]
    
    KNOWN_ROLES = [
        # Core engineering roles
        'systems engineer', 'project manager', 'program manager', 'chief engineer',
        'lead engineer', 'lead systems engineer', 'technical authority',
        'safety engineer', 'quality assurance', 'quality assurance engineer',
        'configuration manager', 'data manager', 'risk manager', 'test engineer',
        'integration engineer', 'verification engineer', 'validation engineer',
        'software engineer', 'hardware engineer', 'design engineer',
        
        # Specialized engineering
        'system safety engineer', 'reliability engineer', 'maintainability engineer',
        'human factors engineer', 'environmental engineer', 'thermal engineer',
        'structural engineer', 'propulsion engineer', 'avionics engineer',
        'manufacturing engineer', 'production engineer', 'logistics engineer',
        'sustaining engineer', 'operations engineer', 'process engineer',
        
        # Leadership/management
        'mission assurance', 'mission assurance manager', 'flight director', 
        'mission director', 'ground controller', 'principal investigator',
        'co-investigator', 'project scientist', 'technical lead',
        'discipline lead', 'functional lead', 'subsystem lead', 'element lead',
        'software lead', 'hardware lead', 'test lead', 'integration lead',
        'verification lead', 'requirements manager', 'test manager',
        
        # Contract/government roles
        'contracting officer', 'contracting officer representative', 'cor',
        'technical monitor', 'government technical representative',
        'authorizing official', 'designated engineering representative',
        'designated airworthiness representative', 'contractor', 'customer',
        'government', 'nasa',
        
        # Teams and groups
        'test team', 'project team', 'technical team', 'development team',
        'integration team', 'review team', 'support team',
        'facility operators', 'shift engineers', 'facility operator',
        'shift engineer', 'test personnel', 'technical staff',
        
        # Boards/panels/groups
        'configuration control board', 'change control board',
        'engineering review board', 'technical review board',
        'system safety panel', 'safety panel', 'review panel', 'review board',
        'independent review team', 'standing review board',
        'interface control working group', 'integrated product team',
        'integrated product team lead', 'working group',
        
        # Executive/directorate
        'mission directorate', 'center director', 'deputy director',
        'associate administrator', 'mission directorate associate administrator',
        'office of the chief engineer', 'project planning and control',
        
        # v3.0.91b: Additional roles for broader document support
        # Agile/Scrum roles
        'scrum master', 'scrum team', 'product owner', 'product manager',
        'sprint team', 'agile team', 'agile coach',
        
        # Executive roles
        'chief innovation officer', 'cino', 'deputy cino', 'deputy pgm',
        'chief architect', 'chief technology officer', 'cto',
        'chief information officer', 'cio', 'chief executive officer', 'ceo',
        'chief operations officer', 'coo', 'chief financial officer', 'cfo',
        
        # General organizational roles
        'stakeholder', 'stakeholders', 'subject matter expert', 'sme',
        'consultant', 'consultant team', 'quality auditor', 'business owner',
        'it pm', 'it project manager', 'information technology project manager',
        'project lead', 'city project manager', 'consultant project lead',
        
        # Support roles
        'sponsor', 'project sponsor', 'executive sponsor',
        'administrator', 'coordinator', 'facilitator',
        
        # v3.0.91c: Additional domain-specific roles
        # IT Security
        'chief information security officer', 'ciso', 'security officer',
        'information security officer', 'cybersecurity analyst',
        
        # Healthcare/Clinical
        'medical monitor', 'medical director', 'clinical director',
        'study coordinator', 'clinical research associate', 'cra',
        'data safety monitoring board', 'institutional review board',
        'ethics committee', 'sponsor medical director'
    ]
    
    ACRONYM_MAP = {
        'pm': 'project manager', 'se': 'systems engineer',
        'lse': 'lead systems engineer', 'ipt': 'integrated product team',
        'cor': 'contracting officer representative',
        'gtr': 'government technical representative',
        'der': 'designated engineering representative',
        'dar': 'designated airworthiness representative',
        'ccb': 'configuration control board', 'erb': 'engineering review board',
        'trb': 'technical review board', 'irt': 'independent review team',
        'srb': 'standing review board', 'oce': 'office of the chief engineer',
        'mdaa': 'mission directorate associate administrator',
        'qa': 'quality assurance', 'ma': 'mission assurance',
        'sma': 'safety and mission assurance', 'ivv': 'independent verification and validation',
        'pp&c': 'project planning and control', 'icwg': 'interface control working group',
        # v3.0.91c additions
        'ciso': 'chief information security officer', 'cra': 'clinical research associate',
        'irb': 'institutional review board', 'dsmb': 'data safety monitoring board'
    }
    
    ACTION_VERBS = {
        'performs': ['perform', 'performs', 'performing', 'execute', 'executes', 'conduct', 'conducts', 'conducting'],
        'approves': ['approve', 'approves', 'approving', 'authorize', 'authorizes', 'sign', 'signs', 'certify', 'certifies'],
        'reviews': ['review', 'reviews', 'reviewing', 'evaluate', 'evaluates', 'assess', 'assesses', 'examine', 'examines'],
        'coordinates': ['coordinate', 'coordinates', 'coordinating', 'collaborate', 'collaborates', 'liaise', 'liaises'],
        'manages': ['manage', 'manages', 'managing', 'oversee', 'oversees', 'direct', 'directs', 'supervise', 'supervises'],
        'supports': ['support', 'supports', 'supporting', 'assist', 'assists', 'help', 'helps', 'aid', 'aids'],
        'verifies': ['verify', 'verifies', 'verifying', 'validate', 'validates', 'confirm', 'confirms', 'check', 'checks'],
        'develops': ['develop', 'develops', 'developing', 'create', 'creates', 'design', 'designs', 'prepare', 'prepares'],
        'maintains': ['maintain', 'maintains', 'maintaining', 'update', 'updates', 'sustain', 'sustains'],
        'ensures': ['ensure', 'ensures', 'ensuring', 'guarantee', 'guarantees', 'assure', 'assures'],
        'provides': ['provide', 'provides', 'providing', 'supply', 'supplies', 'deliver', 'delivers'],
        'receives': ['receive', 'receives', 'receiving', 'obtain', 'obtains'],
        'reports': ['report', 'reports', 'reporting', 'communicate', 'communicates', 'inform', 'informs'],
        'defines': ['define', 'defines', 'defining', 'specify', 'specifies', 'establish', 'establishes'],
        'implements': ['implement', 'implements', 'implementing', 'deploy', 'deploys'],
        'leads': ['lead', 'leads', 'leading'],
        'monitors': ['monitor', 'monitors', 'monitoring', 'track', 'tracks', 'tracking']
    }
    
    FALSE_POSITIVES = [
        # Generic terms
        'the system', 'the document', 'the customer', 'the contractor', 'the government',
        'the agency', 'the organization', 'the project', 'the program',
        'the process', 'the procedure', 'the requirement', 'the specification',
        'the design', 'the product', 'the hardware', 'the software', 'the element',
        'the component', 'the subsystem', 'the interface', 'the data', 'the information',
        'this document', 'this section', 'this chapter', 'this appendix',
        'system', 'document', 'customer', 'contractor', 'government',
        'agency', 'organization', 'project', 'program', 'process', 'procedure',
        
        # Process/discipline names (not roles)
        'systems engineering', 'project management', 'configuration management',
        'risk management', 'data management', 'technical management',
        'quality assurance process', 'verification process', 'validation process',
        
        # Activity-based
        'integration activities', 'verification activities', 'validation activities',
        'test activities', 'review activities', 'design activities', 'development activities',
        'manufacturing activities', 'production activities', 'operations activities',
        
        # Review/document types
        'technical reviews', 'safety reviews', 'design reviews', 'peer reviews',
        'milestone reviews', 'gate reviews', 'phase reviews',
        'interface specifications', 'design specifications', 'requirements specifications',
        'technical requirements', 'functional requirements', 'performance requirements',
        'system requirements', 'software requirements', 'hardware requirements',
        
        # v3.0.91b: Expanded false positives from testing
        # Generic single words that are not roles
        'progress', 'upcoming', 'distinct', 'others', 'addition', 'manner',
        'work', 'test', 'end', 'task', 'plan', 'phase', 'configuration',
        'property', 'resource', 'attachment', 'reports', 'technical',
        'engineering', 'authority', 'international travel', 'coordinating',
        
        # Events and milestones (not roles)
        'test readiness review', 'test readiness', 'test preparation begins',
        'design review', 'preliminary design review', 'critical design review',
        'system requirements review', 'mission readiness review',
        'operational readiness review', 'flight readiness review',
        'phase transition', 'key decision point', 'milestone',
        
        # Facilities and equipment (not roles)
        'panel test facility', 'flight facility', 'test facility',
        'wind tunnel', 'arc jet', 'vacuum chamber', 'clean room',
        'thermal protection', 'support facility', 'support facilities',
        
        # Processes and methodologies
        'reliability centered maintenance', 'condition based maintenance',
        'preventive maintenance', 'corrective maintenance', 'predictive maintenance',
        'property management', 'resource scheduling', 'contract reporting',
        'safety and environmental compliance', 'thermo physics facilities configuration',
        
        # Document and contract elements
        'statement of work', 'scope of work', 'contract data requirements',
        'task order', 'idiq task orders', 'contract line item',
        'operations and maintenance plan', 'technical project',
        
        # Generic phrases that get extracted incorrectly
        'other disciplines', 'other arc organizations', 'foreign systems',
        'all parties', 'project teams', 'wide variety of skills',
        'monthly financial', 'written status report',
        
        # Verb phrases and fragments
        'facilities perform various', 'mission goes beyond just',
        'develop an approach that optimizes', 'facilitate completion of work',
        'operate facilities', 'ensure that utilities', 'provide property',
        'provide the risk', 'coordinate the demand',
        
        # v3.0.91c: Additional false positives from 5-document validation test
        # Organizational disciplines/functions (not roles unless paired with role suffix)
        'mission assurance', 'mission equipment', 'mission systems',
        'configuration control', 'quality control', 'version control',
        'safety and environmental', 'security engineering',
        
        # Truncated or partial role names
        'chief innovation', 'deputy chief', 'assistant deputy',
        
        # IPT/Team fragments
        'staffing integrated product team', 'se ipt lead se',
        'verification engineer'  # Usually too vague without context
    ]
    
    # v3.0.91b: Words that should NEVER be roles on their own
    SINGLE_WORD_EXCLUSIONS = {
        # Generic nouns
        'progress', 'upcoming', 'distinct', 'others', 'addition', 'manner',
        'work', 'test', 'end', 'task', 'plan', 'phase', 'reports', 'property',
        'resource', 'attachment', 'travel', 'training', 'construction',
        'calibration', 'demolition', 'repair', 'installation', 'procurement',
        'schedule', 'budget', 'scope', 'status', 'summary', 'overview',
        'configuration', 'process', 'contract', 'mission', 'officer',
        # Adjectives
        'technical', 'functional', 'operational', 'administrative', 'preliminary',
        'final', 'initial', 'current', 'previous', 'various', 'multiple',
        # Verbs/gerunds
        'coordinating', 'managing', 'performing', 'conducting', 'reviewing',
        'approving', 'verifying', 'validating', 'monitoring', 'tracking',
        'planning', 'scheduling', 'reporting', 'documenting', 'processing',
    }
    
    # Deliverable patterns - items that are work products, not roles
    DELIVERABLE_PATTERNS = [
        # Document types
        r'\b(plan|report|specification|analysis|assessment|study)\b',
        r'\b(document|manual|guide|handbook|procedure|instruction)\b',
        r'\b(drawing|schematic|diagram|model|prototype|mockup)\b',
        r'\b(database|repository|archive|library|registry)\b',
        r'\b(schedule|timeline|roadmap|milestone|gantt)\b',
        r'\b(budget|estimate|proposal|quotation|bid)\b',
        r'\b(contract|agreement|memorandum|charter)\b',
        # Acronyms commonly used for deliverables
        r'\bICD\b|\bSRS\b|\bSDD\b|\bCDRL\b|\bDID\b|\bSOW\b|\bWBS\b',
        r'\bSEMP\b|\bSPMP\b|\bCMP\b|\bQAP\b|\bSAFETY\s+PLAN\b',
        # Test/verification outputs
        r'\btest\s+(report|results|data|log|procedure)\b',
        r'\bverification\s+(report|results|matrix)\b',
        r'\bvalidation\s+(report|results)\b',
    ]
    
    BOUNDARY_WORDS = [
        'and', 'or', 'who', 'which', 'that', 'with', 'for', 'in', 'on', 'at', 'as',
        'to', 'from', 'by', 'is', 'are', 'was', 'were', 'has', 'have', 'had',
        'shall', 'will', 'may', 'can', 'should', 'must', 'could', 'would',
        'the', 'a', 'an', 'this', 'that', 'these', 'those', 'their', 'its',
        'prior', 'before', 'after', 'during', 'while', 'when', 'if', 'unless'
    ]

    def __init__(self, custom_roles: List[str] = None, custom_false_positives: List[str] = None,
                 use_dictionary: bool = True, use_nlp: bool = True):
        """
        Initialize the role extractor with optional customizations.
        
        Args:
            custom_roles: Additional role names to recognize
            custom_false_positives: Additional false positives to ignore
            use_dictionary: If True, attempt to load roles from database dictionary
            use_nlp: If True, use NLP enhancement for better accuracy
        """
        self.known_roles = set(r.lower() for r in self.KNOWN_ROLES)
        if custom_roles:
            self.known_roles.update(r.lower() for r in custom_roles)
        
        # Load from dictionary if available
        if use_dictionary:
            dict_roles = self._load_dictionary_roles()
            if dict_roles:
                self.known_roles.update(r.lower() for r in dict_roles)
        
        self.false_positives = set(fp.lower() for fp in self.FALSE_POSITIVES)
        if custom_false_positives:
            self.false_positives.update(fp.lower() for fp in custom_false_positives)
        
        # Initialize NLP enhancer for better extraction (v3.0.91+)
        self._nlp_enhancer = None
        if use_nlp:
            try:
                from nlp_enhancer import NLPEnhancer
                self._nlp_enhancer = NLPEnhancer()
                _log("NLP enhancer loaded for improved role extraction", level='debug')
            except ImportError:
                _log("NLP enhancer not available, using standard extraction", level='debug')
        
        self._build_patterns()
    
    def _load_dictionary_roles(self) -> List[str]:
        """Load active roles from the database dictionary if available."""
        try:
            import sqlite3
            from pathlib import Path
            
            # Try common database locations
            possible_paths = [
                Path(__file__).parent / 'scan_history.db',
                Path(__file__).parent / 'data' / 'scan_history.db',
                Path.home() / '.twr' / 'scan_history.db'
            ]
            
            for db_path in possible_paths:
                if db_path.exists():
                    conn = sqlite3.connect(str(db_path))
                    cursor = conn.cursor()
                    
                    # Check if table exists
                    cursor.execute("""
                        SELECT name FROM sqlite_master 
                        WHERE type='table' AND name='role_dictionary'
                    """)
                    if cursor.fetchone():
                        cursor.execute("""
                            SELECT role_name, aliases FROM role_dictionary
                            WHERE is_active = 1 AND is_deliverable = 0
                        """)
                        
                        roles = []
                        for row in cursor.fetchall():
                            roles.append(row[0])
                            # Also add aliases
                            if row[1]:
                                try:
                                    import json
                                    aliases = json.loads(row[1])
                                    roles.extend(aliases)
                                except Exception:
                                    pass
                        
                        conn.close()
                        return roles
                    conn.close()
            
            return []
        except Exception as e:
            # Silently fail - dictionary is optional
            return []
    
    def _build_patterns(self):
        """Build compiled regex patterns for role detection."""
        
        # Pattern 1: "The [Role] shall/will/is/has/ensures..."
        self.pattern_the_role = re.compile(
            r'\b[Tt]he\s+([A-Z][a-zA-Z]+(?:\s+[A-Z]?[a-zA-Z]+){0,5}?)\s+'
            r'(?:shall|will|has\s+(?:overall\s+)?responsibility|is\s+(?:skilled|responsible)|'
            r'ensures?|provides?|reviews?|approves?|manages?|coordinates?|performs?|'
            r'maintains?|leads?|oversees?|plays|usually|monitors?)',
            re.MULTILINE
        )
        
        # Pattern 2: "[Role] is responsible for..."
        self.pattern_role_is = re.compile(
            r'\b([A-Z][a-zA-Z]+(?:\s+[A-Z]?[a-zA-Z]+){0,5}?)\s+'
            r'(?:is\s+responsible\s+for|shall\s+\w|will\s+(?:be|ensure|provide|review|approve|manage))',
            re.MULTILINE
        )
        
        # Pattern 3: "by/to/from/with the [Role]"
        self.pattern_by_role = re.compile(
            r'(?:by|to|from|with)\s+the\s+([A-Z][a-zA-Z]+(?:\s+[A-Z]?[a-zA-Z]+){0,4}?)'
            r'(?=\s*[,\.\;]|\s+(?:and|or|who|for|in|on|at|as|prior|before|after)\b)',
            re.MULTILINE
        )
        
        # Pattern 4: "approved/reviewed/verified by [Role]"
        self.pattern_action_by = re.compile(
            r'(?:approved|reviewed|coordinated|submitted|verified|validated|signed|'
            r'certified|authorized|prepared|developed|conducted|performed)\s+'
            r'(?:by|with)\s+(?:the\s+)?([A-Z][a-zA-Z]+(?:\s+[A-Z]?[a-zA-Z]+){0,4}?)'
            r'(?=\s*[,\.\;]|\s+(?:and|or|who|for|in|on|at|as|prior|before|after)\b)',
            re.MULTILINE
        )
        
        # Pattern 5: "[Role]'s responsibility/role/duties"
        self.pattern_possessive = re.compile(
            r"([A-Z][a-zA-Z]+(?:\s+[A-Z]?[a-zA-Z]+){0,4}?)'s\s+"
            r"(?:responsibility|role|duties|function|authority|approval)",
            re.MULTILINE
        )
        
        # Pattern 6: "responsibilities/role of the [Role]"
        self.pattern_responsibilities_of = re.compile(
            r'(?:responsibilities?|roles?|duties|functions?|authority)\s+of\s+(?:the\s+)?'
            r'([A-Z][a-zA-Z]+(?:\s+[A-Z]?[a-zA-Z]+){0,4}?)'
            r'(?=\s*[,\.\;]|\s+(?:and|or|include|are|is)\b)',
            re.MULTILINE
        )
        
        # Pattern 7: Acronym in parentheses "Project Manager (PM)"
        self.pattern_acronym = re.compile(
            r'([A-Z][a-zA-Z]+(?:\s+[A-Z]?[a-zA-Z]+){0,4}?)\s*\(([A-Z][A-Z&]{1,6})\)',
            re.MULTILINE
        )
        
        # Pattern 8: "as the [Role]" / "serve as [Role]"
        self.pattern_as_role = re.compile(
            r'(?:as\s+(?:the\s+)?|serve\s+as\s+(?:the\s+)?|acting\s+(?:as\s+)?(?:the\s+)?)'
            r'([A-Z][a-zA-Z]+(?:\s+[A-Z]?[a-zA-Z]+){0,4}?)'
            r'(?=\s*[,\.\;]|\s+(?:and|or|who|for|in|on|at)\b)',
            re.MULTILINE
        )
        
        # Pattern 9: "notify/inform/contact/report to the [Role]"
        self.pattern_notify = re.compile(
            r'(?:notify|inform|contact|consult|report\s+to)\s+(?:the\s+)?'
            r'([A-Z][a-zA-Z]+(?:\s+[A-Z]?[a-zA-Z]+){0,4}?)'
            r'(?=\s*[,\.\;]|\s+(?:and|or|who|for|in|on|at|prior|before|after)\b)',
            re.MULTILINE
        )
        
        # Pattern 10: Start of sentence "[Role] shall/will/ensures..."
        self.pattern_sentence_start = re.compile(
            r'(?:^|\.\s+)([A-Z][a-zA-Z]+(?:\s+[A-Z]?[a-zA-Z]+){0,4}?)\s+'
            r'(?:shall|will|must|should|is\s+responsible|ensures?|provides?|'
            r'reviews?|approves?|manages?|coordinates?|performs?|maintains?|leads?|oversees?)',
            re.MULTILINE
        )
        
        # Pattern 11: End of sentence "...by the [Role]."
        self.pattern_end_of_sentence = re.compile(
            r'(?:approved|reviewed|verified|validated|signed|authorized|certified)\s+by\s+(?:the\s+)?'
            r'([A-Z][a-zA-Z]+(?:\s+[A-Z]?[a-zA-Z]+){1,4}?)\s*\.',
            re.MULTILINE
        )
        
        # Pattern 12: Role followed by "and" another role (coordinated with X and Y)
        self.pattern_role_and = re.compile(
            r'(?:by|with|to)\s+(?:the\s+)?([A-Z][a-zA-Z]+(?:\s+[A-Z]?[a-zA-Z]+){0,3}?)\s+and\s+'
            r'(?:the\s+)?([A-Z][a-zA-Z]+(?:\s+[A-Z]?[a-zA-Z]+){0,3}?)(?=\s*[,\.\;]|\s+(?:who|for|in|on|at)\b)',
            re.MULTILINE
        )
        
        # Pattern 13: "from the [Role]" / "require approval from the [Role]"
        self.pattern_from_role = re.compile(
            r'(?:from|require(?:s)?\s+(?:approval\s+)?from)\s+(?:the\s+)?'
            r'([A-Z][a-zA-Z]+(?:\s+[A-Z]?[a-zA-Z]+){0,4}?)(?=\s*[,\.\;]|\s+(?:and|or|who|for|in|on|at)\b)',
            re.MULTILINE
        )
        
        # Pattern 14: "the [Role] assists/works with"
        self.pattern_role_assists = re.compile(
            r'\b[Tt]he\s+([A-Z][a-zA-Z]+(?:\s+[A-Z]?[a-zA-Z]+){0,4}?)\s+'
            r'(?:assists?|works?\s+with|supports?|helps?)',
            re.MULTILINE
        )
        
        # Pattern 15: "conducted by the [multi-word] and approved/reviewed by"
        self.pattern_conducted_by = re.compile(
            r'conducted\s+by\s+(?:the\s+)?([A-Z][a-zA-Z]+(?:\s+[A-Z]?[a-zA-Z]+){0,4}?)\s+and',
            re.MULTILINE
        )
        
        # Pattern 16: "analyzed/assembles/processes" etc - past tense verbs followed by roles
        self.pattern_role_action_object = re.compile(
            r'\b[Tt]he\s+([A-Z][a-zA-Z]+(?:\s+[A-Z]?[a-zA-Z]+){0,3}?)\s+'
            r'(?:analyzes?|assembles?|processes?|develops?|creates?|builds?|designs?|tests?|integrates?)',
            re.MULTILINE
        )
        
        # Pattern 17: "the [Role] in managing/in doing..."
        self.pattern_role_in = re.compile(
            r'\b(?:the\s+)?([A-Z][a-zA-Z]+(?:\s+[A-Z]?[a-zA-Z]+){0,4}?)\s+in\s+managing',
            re.MULTILINE
        )
        
        # Pattern 18: "require[s] approval from the [Role] and [Role]"
        self.pattern_approval_from = re.compile(
            r'(?:require|requires)\s+(?:approval\s+)?from\s+(?:the\s+)?'
            r'([A-Z][a-zA-Z]+(?:\s+[A-Z]?[a-zA-Z]+){0,3}?)\s+and\s+(?:the\s+)?'
            r'([A-Z][a-zA-Z]+(?:\s+[A-Z]?[a-zA-Z]+){0,3}?)(?=\s*[,\.\;])',
            re.MULTILINE
        )
        
        # Pattern 19: Capture "lead [role]" patterns specifically
        self.pattern_lead_role = re.compile(
            r'\b(?:The\s+)?([Ll]ead\s+[A-Z]?[a-zA-Z]+(?:\s+[A-Z]?[a-zA-Z]+){0,2}?)\s+'
            r'(?:shall|will|ensures?|provides?|is\s+responsible|has)',
            re.MULTILINE
        )
        
        # Pattern 20: "performed by the [role] and" - capture compound patterns
        self.pattern_performed_by_and = re.compile(
            r'(?:performed|conducted|reviewed|approved|verified)\s+by\s+(?:the\s+)?'
            r'([A-Z][a-zA-Z]+(?:\s+[A-Z]?[a-zA-Z]+){0,3}?)\s+and\s+(?:the\s+|approved\s+by\s+(?:the\s+)?)?'
            r'([A-Z][a-zA-Z]+(?:\s+[A-Z]?[a-zA-Z]+){0,3}?)(?=\s*[,\.\;])',
            re.MULTILINE
        )
        
        # Pattern 21: "by the [multi-word-role]." at sentence end with multi-word support
        self.pattern_by_multi_word = re.compile(
            r'by\s+(?:the\s+)?([A-Z][a-zA-Z]+(?:\s+[A-Z][a-zA-Z]+){1,4}?)\s*\.',
            re.MULTILINE
        )
        
        # Pattern 22: "as performed by the [role]"
        self.pattern_as_performed = re.compile(
            r'as\s+performed\s+by\s+(?:the\s+)?([A-Z][a-zA-Z]+(?:\s+[A-Z]?[a-zA-Z]+){0,4}?)(?=\s+and)',
            re.MULTILINE
        )
        
        # Pattern 23: "The [role] supports" / "The [role] assists"
        self.pattern_role_supports = re.compile(
            r'\b[Tt]he\s+([a-zA-Z]+(?:\s+[a-zA-Z]+){0,3}?)\s+(?:supports?|assists?)\s+(?:the\s+)?',
            re.MULTILINE
        )
        
        # Pattern 24: "approval from the [Role] and the [Role]" 
        self.pattern_approval_chain = re.compile(
            r'approval\s+from\s+(?:the\s+)?([A-Z][a-zA-Z]+(?:\s+[A-Z]?[a-zA-Z]+){0,3}?)\s+and\s+(?:the\s+)?([A-Z][a-zA-Z]+(?:\s+[A-Z]?[a-zA-Z]+){0,3}?)(?=\s*[,\.\;])',
            re.MULTILINE
        )

    def _clean_candidate(self, candidate: str) -> str:
        """Clean and normalize a candidate role string."""
        candidate = candidate.strip()
        
        # Reject candidates with weird spacing patterns (OCR/parsing errors)
        # E.g., "Fa Ci Lities" instead of "Facilities"
        words = candidate.split()
        
        # Check for many single-letter words or 2-letter words in sequence (OCR error)
        short_word_count = sum(1 for w in words if len(w) <= 2 and w.isalpha())
        if short_word_count >= 3 and len(words) >= 4:
            return ""  # Likely an OCR error like "Fa Ci Lities Pro Ject"
        
        # Check for alternating case pattern that suggests broken text
        if len(words) >= 3:
            all_title_case = all(w[0].isupper() and (len(w) == 1 or w[1:].islower()) for w in words if w.isalpha())
            avg_word_len = sum(len(w) for w in words) / len(words)
            if all_title_case and avg_word_len < 4:
                return ""  # Likely "Fa Ci Li Ties" pattern
        
        # Remove trailing boundary words (conjunctions, prepositions, etc.)
        while words and words[-1].lower() in self.BOUNDARY_WORDS:
            words.pop()
        
        # Remove leading boundary words
        while words and words[0].lower() in self.BOUNDARY_WORDS:
            words.pop(0)
        
        # Remove leading articles
        while words and words[0].lower() in ['the', 'a', 'an']:
            words.pop(0)
        
        # Extended list of verbs that shouldn't be part of role names
        # v3.0.91b: Removed 'leads' as it can be part of role names like "Technical Leads"
        trailing_verbs = [
            # Action verbs
            'analyzes', 'assembles', 'processes', 'develops', 'creates',
            'builds', 'designs', 'tests', 'integrates', 'works', 'assists',
            'monitors', 'reviews', 'approves', 'verifies', 'validates',
            'analyzed', 'assembled', 'processed', 'developed', 'created',
            'performs', 'coordinates', 'manages', 'ensures', 'provides',
            'maintains', 'oversees', 'supports', 'implements',
            # Gerunds
            'analyzing', 'assembling', 'processing', 'developing', 'creating',
            'building', 'designing', 'testing', 'integrating', 'working',
            'assisting', 'monitoring', 'reviewing', 'approving', 'verifying',
            # Object words that indicate overextended extraction
            'system', 'reliability', 'components', 'data', 'requirements',
            'specifications', 'documents', 'activities', 'functions',
            'interfaces', 'design', 'analysis', 'report', 'documentation'
        ]
        
        # v3.0.91b: Don't strip 'leads' if preceded by a role modifier (it's part of role name)
        # "Technical Leads" should stay intact, but "The engineer leads" should lose "leads"
        if words and words[-1].lower() == 'leads' and len(words) >= 2:
            # Check if previous word suggests this is a role name
            role_modifiers = {'technical', 'team', 'project', 'program', 'system', 'software',
                             'hardware', 'test', 'qa', 'quality', 'safety', 'design', 'integration'}
            if words[-2].lower() not in role_modifiers:
                words.pop()  # Remove 'leads' if it's likely a verb
        
        while words and words[-1].lower() in trailing_verbs:
            words.pop()
        
        # Also remove leading verbs (shouldn't start a role name)
        leading_verbs = [
            'shall', 'will', 'may', 'can', 'should', 'must',
            'performs', 'coordinates', 'manages', 'ensures', 'provides',
            'reviews', 'approves', 'verifies', 'validates', 'monitors'
        ]
        while words and words[0].lower() in leading_verbs:
            words.pop(0)
        
        # If result is too long or contains verb phrases, try to truncate
        result = ' '.join(words)
        
        # Check for known patterns where extraction went too far
        verb_phrases = ['analyzes system', 'assembles components', 'processes data',
                       'develops requirements', 'creates documents', 'performs analysis',
                       'conducts review', 'provides support', 'manages project']
        for vp in verb_phrases:
            if vp in result.lower():
                # Truncate at the verb
                verb = vp.split()[0]
                idx = result.lower().find(verb)
                if idx > 0:
                    result = result[:idx].strip()
        
        return result

    def _is_valid_role(self, candidate: str) -> Tuple[bool, float]:
        """Determine if a candidate string is likely a valid role."""
        candidate = self._clean_candidate(candidate)
        candidate_lower = candidate.lower().strip()
        
        # v3.0.91b: Check for known acronyms FIRST (before length check)
        # These are short but valid role identifiers
        valid_acronyms = {'cor', 'pm', 'se', 'lse', 'ipt', 'ccb', 'erb', 'ta', 'pi', 
                         'qa', 'qae', 'cm', 'sme', 'der', 'dar', 'gtr', 'co', 'cotr',
                         # Executive acronyms
                         'cino', 'cto', 'cio', 'ceo', 'coo', 'cfo', 'pgm', 'dpgm',
                         # IT roles
                         'dba', 'sa', 'sysadmin'}
        if candidate_lower in valid_acronyms:
            return True, 0.92
        
        # Basic checks
        if len(candidate_lower) < 4 or len(candidate.split()) < 1:
            return False, 0.0
        
        if len(candidate.split()) > 6:
            return False, 0.0
        
        # v3.0.91b: Reject single words that are never roles
        words = candidate_lower.split()
        if len(words) == 1 and candidate_lower in self.SINGLE_WORD_EXCLUSIONS:
            return False, 0.0
        
        # v3.0.91b: Reject phrases starting with noise words
        noise_starters = {
            'the', 'a', 'an', 'this', 'that', 'all', 'any', 'each', 'some',
            'contract', 'provide', 'ensure', 'facilitate', 'complete',
            'develop', 'perform', 'operate', 'coordinate', 'manage',
            'services', 'requirements', 'process', 'phase', 'work',
            'idiq', 'task', 'data', 'report', 'plan', 'end', 'wide',
            'construction', 'calibration', 'demolition', 'repair',
            'manner', 'range', 'monthly', 'foreign', 'addition',
            'written', 'final', 'initial', 'current', 'other', 'various',
            # v3.0.91b: Additional noise starters
            'responsible', 'accountable', 'serves', 'serving', 'acts',
            'acting', 'works', 'working', 'reports', 'reporting',
            'directly', 'overall', 'primary', 'secondary', 'main'
        }
        if words and words[0] in noise_starters:
            return False, 0.0
        
        # v3.0.91b: Reject phrases containing connector words in positions 2-4
        connector_words = {'is', 'are', 'was', 'were', 'be', 'being', 'been',
                          'shall', 'will', 'must', 'may', 'can', 'could',
                          'that', 'which', 'who', 'whom', 'whose',
                          'so', 'such', 'very', 'just', 'only', 'also',
                          'including', 'excluding', 'regarding', 'concerning',
                          'for'}  # Added 'for' to catch "Responsible for..."
        if len(words) > 1:
            for word in words[1:4]:  # Check words 2-4
                if word in connector_words:
                    return False, 0.0
        
        # v3.0.91b: Reject phrases with "and" followed by another role word
        # e.g., "COR And Contracting Officer" should be split, not one role
        if 'and' in words:
            and_idx = words.index('and')
            # If there's a role-like word after "and", this is likely two roles combined
            if and_idx > 0 and and_idx < len(words) - 1:
                after_and = ' '.join(words[and_idx+1:])
                # Check if the part after "and" looks like a role
                role_suffixes = {'manager', 'engineer', 'officer', 'director', 'lead', 
                                'analyst', 'team', 'coordinator', 'specialist'}
                if any(s in after_and for s in role_suffixes):
                    return False, 0.0
        
        # v3.0.91b: Reject "X The Y" patterns (sentence fragments)
        if len(words) >= 3 and 'the' in words[1:]:
            the_idx = words.index('the') if 'the' in words else -1
            if the_idx > 0:  # "the" is not at start
                return False, 0.0
        
        # v3.0.91b: Reject phrases ending with noise words
        noise_endings = {'begins', 'ends', 'various', 'just', 'only',
                        'personnel', 'activities', 'services', 'requirements',
                        'orders', 'systems', 'facilities', 'resources'}
        if words and words[-1] in noise_endings:
            return False, 0.0
        
        # v3.0.91d: Check EXPLICIT false positives FIRST - overrides everything
        # This ensures items like "Verification Engineer" or "Mission Assurance" are filtered
        # even if they appear in known_roles
        if candidate_lower in self.false_positives:
            return False, 0.0
        
        # v3.0.12b: Check for known roles (high confidence)
        # This prevents false-positive rules from rejecting real roles like "Program Manager"
        if candidate_lower in self.known_roles:
            return True, 0.95
        
        # v3.0.12b: Check for role suffix BEFORE generic false-positive patterns
        # Strong role suffixes override generic false-positive patterns
        strong_role_suffixes = ['engineer', 'manager', 'lead', 'analyst', 'specialist', 
                               'coordinator', 'owner', 'reviewer', 'approver', 'author',
                               'director', 'officer', 'supervisor', 'architect']
        for suffix in strong_role_suffixes:
            if candidate_lower.endswith(suffix):
                prefix = candidate_lower[:-len(suffix)].strip()
                if prefix:  # Has a modifier before the suffix
                    return True, 0.90
        
        # Check against partial false positive matches
        for fp in self.false_positives:
            if candidate_lower.startswith(fp + ' ') or candidate_lower.endswith(' ' + fp):
                return False, 0.0
        
        # Check for problematic patterns
        bad_patterns = ['reviewed by', 'approved by', 'conducted by', 'coordinated with',
                       'performed by', 'verified by', 'validated by', 'assists the',
                       'works with the', 'in managing', 'goes beyond', 'perform various',
                       'that optimizes', 'completion of', 'demand for']
        if any(bp in candidate_lower for bp in bad_patterns):
            return False, 0.0
        
        # Check for fragments that don't make sense as roles
        if candidate_lower.startswith('manager assists') or candidate_lower.startswith('engineer analyzes'):
            return False, 0.0
        
        # Single word "Lead" or "Manager" alone is too generic unless it's a known role
        if candidate_lower in ['lead', 'manager', 'engineer', 'director', 'analyst']:
            return False, 0.0
        
        # Check for activity endings
        activity_endings = ['activities', 'reviews', 'specifications', 'requirements', 
                          'procedures', 'processes', 'tasks', 'efforts', 'orders',
                          'travel', 'maintenance', 'protection', 'facility']
        for ending in activity_endings:
            if candidate_lower.endswith(ending):
                # Exception: allow if it's clearly a role
                if not any(s in candidate_lower for s in ['engineer', 'manager', 'lead', 'officer']):
                    return False, 0.0
        
        # Check for role suffixes (general)
        for suffix in self.ROLE_SUFFIXES:
            if candidate_lower.endswith(suffix):
                prefix = candidate_lower[:-len(suffix)].strip()
                if prefix:
                    has_valid_modifier = any(
                        word.lower() in [m.lower() for m in self.ROLE_MODIFIERS]
                        for word in prefix.split()
                    )
                    if has_valid_modifier:
                        return True, 0.90
                    elif len(prefix.split()) <= 2:
                        return True, 0.75
                    else:
                        return True, 0.60
                else:
                    return True, 0.50
        
        # Check for role suffix in middle of phrase
        for word in words:
            if word in self.ROLE_SUFFIXES:
                return True, 0.65
        
        # Check for modifier + something pattern
        if len(words) >= 2 and words[0] in [m.lower() for m in self.ROLE_MODIFIERS]:
            return True, 0.55
        
        # Check for partial match with known roles
        for known in self.known_roles:
            if known in candidate_lower and len(candidate_lower) < len(known) + 15:
                return True, 0.70
        
        return False, 0.0
    
    def _is_deliverable(self, candidate: str) -> bool:
        """Check if a candidate is likely a deliverable/work product rather than a role."""
        candidate_lower = candidate.lower()
        
        # Check against deliverable patterns
        for pattern in self.DELIVERABLE_PATTERNS:
            if re.search(pattern, candidate_lower, re.IGNORECASE):
                return True
        
        # Check for common deliverable naming patterns
        deliverable_indicators = [
            'plan', 'report', 'specification', 'document', 'manual',
            'drawing', 'schematic', 'schedule', 'budget', 'proposal',
            'matrix', 'database', 'archive', 'log', 'checklist'
        ]
        
        words = candidate_lower.split()
        if words and words[-1] in deliverable_indicators:
            return True
        
        return False
    
    def classify_extraction(self, candidate: str) -> dict:
        """
        Classify an extraction as role, deliverable, or unknown.
        Returns dict with type, confidence, and reasoning.
        
        v3.0.12b: Role suffix wins tie-break. If candidate ends with a strong
        role suffix (engineer, manager, lead, etc.), classify as role even if
        deliverable keywords are present.
        
        v3.0.105: BUG-004 FIX - Strong role suffix now ALWAYS wins tiebreak,
        regardless of _is_valid_role result. This ensures "Report Engineer"
        is classified as role, not deliverable.
        """
        candidate_lower = candidate.lower().strip()
        words = candidate_lower.split()
        
        # v3.0.12b: Strong role suffixes that win tie-break against deliverable keywords
        strong_role_suffixes = {'engineer', 'manager', 'lead', 'analyst', 'specialist', 
                               'coordinator', 'owner', 'reviewer', 'approver', 'author',
                               'director', 'officer', 'supervisor', 'architect', 'sme',
                               'integrator', 'administrator', 'technician', 'inspector'}
        
        # Check if ends with strong role suffix - this wins tie-break
        has_strong_role_suffix = words and words[-1] in strong_role_suffixes
        
        # v3.0.105 FIX: If has strong role suffix, classify as role IMMEDIATELY
        # Don't rely on _is_valid_role which may have false negatives
        if has_strong_role_suffix:
            # Still try to get confidence from _is_valid_role if possible
            is_valid, confidence = self._is_valid_role(candidate)
            # Use higher confidence if valid, otherwise use base 0.9 for suffix match
            effective_confidence = max(confidence, 0.9) if is_valid else 0.9
            return {
                'type': 'role',
                'confidence': effective_confidence,
                'reason': f'Strong role suffix ({words[-1]}) wins tie-break'
            }
        
        # Check if it's a deliverable (only if no strong role suffix)
        if self._is_deliverable(candidate):
            return {
                'type': 'deliverable',
                'confidence': 0.85,
                'reason': 'Matches deliverable pattern'
            }
        
        # Check if it's a valid role
        is_valid, confidence = self._is_valid_role(candidate)
        if is_valid:
            return {
                'type': 'role',
                'confidence': confidence,
                'reason': 'Matches role pattern'
            }
        
        # Check if it's in false positives
        if candidate_lower in self.false_positives:
            return {
                'type': 'false_positive',
                'confidence': 0.9,
                'reason': 'Known false positive'
            }
        
        return {
            'type': 'unknown',
            'confidence': 0.3,
            'reason': 'Could not classify'
        }
    
    def _create_extracted_role(self, canonical_name: str, original_text: str = None) -> ExtractedRole:
        """Create an ExtractedRole with entity classification.
        
        v3.0.12: All role creation goes through this method to ensure
        entity_kind is always populated.
        """
        text_to_classify = original_text or canonical_name
        classification = self.classify_extraction(text_to_classify)
        
        # Map classification type to EntityKind enum
        kind_map = {
            'role': EntityKind.ROLE,
            'deliverable': EntityKind.DELIVERABLE,
            'false_positive': EntityKind.UNKNOWN,
            'unknown': EntityKind.UNKNOWN
        }
        entity_kind = kind_map.get(classification['type'], EntityKind.UNKNOWN)
        
        return ExtractedRole(
            canonical_name=canonical_name,
            entity_kind=entity_kind,
            kind_confidence=classification['confidence'],
            kind_reason=classification['reason']
        )
    
    def _extract_responsibility(self, text: str, role_match: str) -> Tuple[str, str]:
        """Extract the responsibility associated with a role from context."""
        text_lower = text.lower()
        role_lower = role_match.lower()
        
        role_pos = text_lower.find(role_lower)
        if role_pos == -1:
            return "", "unknown"
        
        after_role = text[role_pos + len(role_match):].strip()
        action_type = "unknown"
        responsibility = ""
        
        for atype, verbs in self.ACTION_VERBS.items():
            for verb in verbs:
                verb_match = re.search(rf'\b{verb}\b', after_role[:100], re.IGNORECASE)
                if verb_match:
                    action_type = atype
                    verb_pos = verb_match.end()
                    remainder = after_role[verb_pos:].strip()
                    end_match = re.search(r'[\.;]', remainder)
                    if end_match:
                        responsibility = remainder[:end_match.start()].strip()
                    else:
                        responsibility = remainder[:100].strip()
                    break
            if action_type != "unknown":
                break
        
        if action_type == "unknown":
            resp_match = re.search(r'is\s+responsible\s+for\s+(.{10,100}?)(?:\.|;|$)', after_role, re.IGNORECASE)
            if resp_match:
                action_type = "performs"
                responsibility = resp_match.group(1).strip()
        
        return responsibility, action_type
    
    def _normalize_role(self, role: str) -> str:
        """Normalize a role name to a canonical form."""
        role = self._clean_candidate(role)
        normalized = ' '.join(role.split()).title()
        
        # Keep abbreviations uppercase
        abbreviations = {
            'Pm': 'PM', 'Se': 'SE', 'Lse': 'LSE', 'Ipt': 'IPT', 'Cor': 'COR',
            'Gtr': 'GTR', 'Der': 'DER', 'Dar': 'DAR', 'Ccb': 'CCB', 'Erb': 'ERB',
            'Trb': 'TRB', 'Irt': 'IRT', 'Srb': 'SRB', 'Oce': 'OCE', 'Mdaa': 'MDAA',
            'Qa': 'QA', 'Ma': 'MA', 'Sma': 'SMA', 'It': 'IT', 'Iv&V': 'IV&V',
            'Ivv': 'IVV', 'Pp&C': 'PP&C', 'Icwg': 'ICWG'
        }
        
        for abbr, correct in abbreviations.items():
            normalized = re.sub(rf'\b{abbr}\b', correct, normalized)
        
        return normalized
    
    # Role aliases for merging similar roles
    ROLE_ALIASES = {
        # Singular/plural and minor variations
        'Systems Engineer': ['System Engineer', 'Systems Engineers', 'System Engineers', 'Sys Engineer'],
        'Software Engineer': ['Software Engineers', 'SW Engineer', 'SW Engineers'],
        'Hardware Engineer': ['Hardware Engineers', 'HW Engineer', 'HW Engineers'],
        'Quality Assurance': ['Quality Assurance Engineer', 'QA Engineer', 'QA', 'Quality Engineer'],
        'Project Manager': ['Project Managers', 'Proj Manager', 'PM'],
        'Program Manager': ['Program Managers', 'Prog Manager'],
        'Test Engineer': ['Test Engineers', 'Testing Engineer', 'Test Eng'],
        'Safety Engineer': ['Safety Engineers', 'System Safety Engineer', 'Systems Safety Engineer'],
        'Reliability Engineer': ['Reliability Engineers', 'Rel Engineer'],
        'Integration Engineer': ['Integration Engineers', 'Integrator', 'System Integrator'],
        'Configuration Manager': ['Configuration Management', 'CM', 'Config Manager'],
        'Technical Lead': ['Tech Lead', 'Technical Leads', 'Tech Leads'],
        'Chief Engineer': ['Chief Engineers', 'CE', 'Chief Eng'],
        'Lead Engineer': ['Lead Engineers', 'Lead Eng'],
        'Requirements Engineer': ['Requirements Engineers', 'Req Engineer', 'Requirements Analyst'],
        'Design Engineer': ['Design Engineers', 'Designer'],
        'Verification Engineer': ['Verification Engineers', 'V&V Engineer', 'Verification'],
        'Validation Engineer': ['Validation Engineers', 'Validation'],
        'Subcontractor': ['Subcontractors', 'Sub-Contractor', 'Sub Contractor'],
        'Contractor': ['Contractors', 'Prime Contractor'],
        'Customer': ['Customers', 'Client', 'End User'],
        'Government': ['Government Representative', 'Govt', 'Government Customer'],
    }
    
    def _get_canonical_role(self, role_name: str) -> Tuple[str, str]:
        """
        Get the canonical name for a role, merging aliases.
        Returns (canonical_name, original_variant).
        """
        normalized = self._normalize_role(role_name)
        normalized_lower = normalized.lower()
        
        # Check if this role matches any alias
        for canonical, aliases in self.ROLE_ALIASES.items():
            if normalized_lower == canonical.lower():
                return canonical, normalized
            for alias in aliases:
                if normalized_lower == alias.lower():
                    return canonical, normalized
        
        # Check for partial matches (e.g., "System Engineer" should match "Systems Engineer")
        for canonical, aliases in self.ROLE_ALIASES.items():
            # Check similarity - simple approach: compare without 's' suffix
            canon_base = canonical.lower().replace('systems', 'system').replace('engineers', 'engineer')
            norm_base = normalized_lower.replace('systems', 'system').replace('engineers', 'engineer')
            if canon_base == norm_base:
                return canonical, normalized
        
        return normalized, normalized
    
    def _get_sentence_context(self, text: str, match_start: int, match_end: int) -> str:
        """Extract the sentence containing the match."""
        sentence_start = max(0, text.rfind('.', 0, match_start) + 1)
        sentence_end = text.find('.', match_end)
        if sentence_end == -1:
            sentence_end = len(text)
        else:
            sentence_end += 1
        
        context = text[sentence_start:sentence_end].strip()
        return ' '.join(context.split())
    
    def extract_from_text(self, text: str, source_location: str = "unknown") -> Dict[str, ExtractedRole]:
        """Extract roles from plain text."""
        extracted_roles: Dict[str, ExtractedRole] = {}
        seen_matches = set()
        
        patterns = [
            ('the_role', self.pattern_the_role, False),
            ('role_is', self.pattern_role_is, False),
            ('by_role', self.pattern_by_role, False),
            ('action_by', self.pattern_action_by, False),
            ('possessive', self.pattern_possessive, False),
            ('responsibilities_of', self.pattern_responsibilities_of, False),
            ('acronym', self.pattern_acronym, True),
            ('as_role', self.pattern_as_role, False),
            ('notify', self.pattern_notify, False),
            ('sentence_start', self.pattern_sentence_start, False),
            ('end_of_sentence', self.pattern_end_of_sentence, False),
            ('role_and', self.pattern_role_and, True),
            ('from_role', self.pattern_from_role, False),
            ('role_assists', self.pattern_role_assists, False),
            ('conducted_by', self.pattern_conducted_by, False),
            ('role_action_object', self.pattern_role_action_object, False),
            ('role_in', self.pattern_role_in, False),
            ('approval_from', self.pattern_approval_from, True),
            ('lead_role', self.pattern_lead_role, False),
            ('performed_by_and', self.pattern_performed_by_and, True),
            ('by_multi_word', self.pattern_by_multi_word, False),
            ('as_performed', self.pattern_as_performed, False),
            ('role_supports', self.pattern_role_supports, False),
            ('approval_chain', self.pattern_approval_chain, True)
        ]
        
        for pattern_name, pattern, has_multiple_groups in patterns:
            for match in pattern.finditer(text):
                candidates = []
                
                if pattern_name == 'acronym':
                    candidates.append((match.group(1), match.group(2)))
                elif pattern_name == 'role_and':
                    candidates.append((match.group(1), None))
                    candidates.append((match.group(2), None))
                else:
                    candidates.append((match.group(1), None))
                
                for candidate, acronym in candidates:
                    candidate = self._clean_candidate(candidate)
                    
                    if not candidate:
                        continue
                    
                    is_valid, confidence = self._is_valid_role(candidate)
                    
                    if not is_valid:
                        continue
                    
                    # Get canonical name and original variant
                    canonical, variant = self._get_canonical_role(candidate)
                    
                    match_key = (canonical, match.start())
                    if match_key in seen_matches:
                        continue
                    seen_matches.add(match_key)
                    
                    context = self._get_sentence_context(text, match.start(), match.end())
                    responsibility, action_type = self._extract_responsibility(context, candidate)
                    
                    occurrence = RoleOccurrence(
                        role=candidate,
                        context=context,
                        responsibility=responsibility,
                        action_type=action_type,
                        location=source_location,
                        confidence=confidence
                    )
                    
                    if canonical not in extracted_roles:
                        # v3.0.12: Use helper to ensure entity_kind is populated
                        extracted_roles[canonical] = self._create_extracted_role(canonical, candidate)
                    
                    role_entry = extracted_roles[canonical]
                    role_entry.variants.add(variant)
                    if variant != canonical:
                        role_entry.variants.add(candidate)  # Also add original text
                    role_entry.occurrences.append(occurrence)
                    
                    if responsibility:
                        role_entry.responsibilities.append(responsibility)
                    
                    role_entry.action_types[action_type] += 1
                    
                    if acronym:
                        role_entry.variants.add(acronym)
        
        # ADDITIONAL PASS: Direct scan for known roles that might have been missed
        # This catches cases like "Systems Engineer" in unusual sentence structures
        extracted_roles = self._scan_for_known_roles(text, extracted_roles, seen_matches, source_location)
        
        # v2.9.2 A9: Parse formal "Responsibilities" sections for enhanced role extraction
        extracted_roles = self._parse_responsibilities_sections(text, extracted_roles, source_location)
        
        # v3.0.91+: NLP Enhancement pass for additional role detection
        if self._nlp_enhancer:
            extracted_roles = self._apply_nlp_enhancement(text, extracted_roles, seen_matches, source_location)
        
        return extracted_roles
    
    def _apply_nlp_enhancement(self, text: str, existing_roles: Dict[str, ExtractedRole],
                               seen_matches: set, source_location: str) -> Dict[str, ExtractedRole]:
        """
        Apply NLP enhancement for additional role detection (v3.0.91+).
        
        Uses pattern matching and sklearn clustering for:
        - Additional role detection via enhanced patterns
        - Role clustering for deduplication
        - Confidence adjustment based on context
        """
        if not self._nlp_enhancer:
            return existing_roles
        
        try:
            # Extract additional roles using NLP patterns
            nlp_entities = self._nlp_enhancer.extract_roles_enhanced(text, source_location)
            
            for entity in nlp_entities:
                if entity.entity_type not in ('ROLE', 'ORG'):
                    continue
                
                # Skip if already found
                canonical = entity.text.lower()
                if canonical in [r.lower() for r in existing_roles.keys()]:
                    continue
                
                # v3.0.91b: Use full _is_valid_role validation
                is_valid, confidence = self._is_valid_role(entity.text)
                if not is_valid:
                    continue
                
                # Use the higher of NLP confidence or validation confidence
                final_confidence = max(entity.confidence, confidence)
                
                # Create new role entry
                if final_confidence >= 0.60:  # Threshold for NLP-detected roles
                    canonical_name, variant = self._get_canonical_role(entity.text)
                    
                    if canonical_name not in existing_roles:
                        existing_roles[canonical_name] = self._create_extracted_role(
                            canonical_name, entity.text
                        )
                    
                    role_entry = existing_roles[canonical_name]
                    role_entry.variants.add(entity.text)
                    
                    occurrence = RoleOccurrence(
                        role=entity.text,
                        context=entity.context,
                        responsibility="",
                        action_type="nlp_detected",
                        location=source_location,
                        confidence=final_confidence
                    )
                    role_entry.occurrences.append(occurrence)
                    role_entry.action_types['nlp_detected'] += 1
            
            _log(f"NLP enhancement found {len(nlp_entities)} additional entities", level='debug')
            
        except Exception as e:
            _log(f"NLP enhancement failed: {e}", level='debug')
        
        return existing_roles
    
    def _parse_responsibilities_sections(self, text: str, existing_roles: Dict[str, 'ExtractedRole'],
                                         source_location: str) -> Dict[str, 'ExtractedRole']:
        """
        v2.9.2 A9: Parse formal "Responsibilities" sections that list role duties.
        
        Handles formats like:
        - "RESPONSIBILITIES" / "Roles and Responsibilities" section headers
        - "[Role Name]:" followed by bullet points or numbered items
        - Tabular formats with Role | Responsibility columns
        """
        
        # Pattern to find responsibility section headers
        section_patterns = [
            # Standard section headers
            re.compile(r'(?:^|\n)\s*(?:\d+\.?\s*)?(?:ROLES?\s+AND\s+)?RESPONSIBILITIES\s*(?:\n|$)', re.IGNORECASE | re.MULTILINE),
            re.compile(r'(?:^|\n)\s*(?:\d+\.?\s*)?ORGANIZATIONAL\s+RESPONSIBILITIES\s*(?:\n|$)', re.IGNORECASE | re.MULTILINE),
            re.compile(r'(?:^|\n)\s*(?:\d+\.?\s*)?RESPONSIBILITY\s+MATRIX\s*(?:\n|$)', re.IGNORECASE | re.MULTILINE),
        ]
        
        # Find start of responsibilities section
        section_start = None
        for pattern in section_patterns:
            match = pattern.search(text)
            if match:
                section_start = match.end()
                break
        
        if section_start is None:
            return existing_roles
        
        # Find end of section (next major heading or end of text)
        section_end_pattern = re.compile(r'\n\s*(?:\d+\.?\s*)?[A-Z][A-Z\s]{5,}(?:\n|$)')
        end_match = section_end_pattern.search(text, section_start + 100)  # Skip at least 100 chars
        section_end = end_match.start() if end_match else len(text)
        
        section_text = text[section_start:section_end]
        
        # Pattern for role followed by responsibilities
        # Format: "Role Name:" or "Role Name -" followed by text
        role_duty_pattern = re.compile(
            r'(?:^|\n)\s*(?:•|\*|[-–—]|\d+[.\)]\s*)?'  # Optional bullet/number
            r'([A-Z][A-Za-z]+(?:\s+[A-Z]?[a-zA-Z]+){0,4})'  # Role name (1-5 words, starts with caps)
            r'\s*(?::|[-–—]|shall|will|is responsible)\s*'  # Separator
            r'(.{10,200}?)(?=\n|$)',  # Responsibility text (10-200 chars)
            re.MULTILINE
        )
        
        for match in role_duty_pattern.finditer(section_text):
            role_text = match.group(1).strip()
            responsibility = match.group(2).strip()
            
            # Validate this is actually a role
            role_text = self._clean_candidate(role_text)
            if not role_text:
                continue
            
            is_valid, confidence = self._is_valid_role(role_text)
            if not is_valid:
                continue
            
            canonical, variant = self._get_canonical_role(role_text)
            
            # Create or update role entry
            if canonical not in existing_roles:
                # v3.0.12: Use helper to ensure entity_kind is populated
                existing_roles[canonical] = self._create_extracted_role(canonical, role_text)
            
            role_entry = existing_roles[canonical]
            role_entry.variants.add(variant)
            
            # Add responsibility if not duplicate
            if responsibility and responsibility not in role_entry.responsibilities:
                role_entry.responsibilities.append(responsibility)
            
            # Create occurrence record
            occurrence = RoleOccurrence(
                role=role_text,
                context=match.group(0).strip()[:150],
                responsibility=responsibility,
                action_type='performs',  # Default for responsibility assignments
                location=f"{source_location} (Responsibilities Section)",
                confidence=0.85  # High confidence from formal section
            )
            role_entry.occurrences.append(occurrence)
            role_entry.action_types['performs'] += 1
        
        return existing_roles
    
    def _scan_for_known_roles(self, text: str, existing_roles: Dict[str, 'ExtractedRole'], 
                              seen_matches: set, source_location: str) -> Dict[str, 'ExtractedRole']:
        """
        Directly scan for known roles that may have been missed by pattern matching.
        This ensures high-confidence roles like "Systems Engineer" are always found.
        """
        text_lower = text.lower()
        
        # Combine known roles with their aliases for comprehensive matching
        all_known = set(self.known_roles)
        for canonical, aliases in self.ROLE_ALIASES.items():
            all_known.add(canonical.lower())
            for alias in aliases:
                all_known.add(alias.lower())
        
        for known_role in all_known:
            if len(known_role) < 5:  # Skip short acronyms to avoid false matches
                continue
            
            # v3.0.91d: Skip if this role is in false_positives
            if known_role in self.false_positives:
                continue
            
            # Find all occurrences of this known role
            start = 0
            while True:
                pos = text_lower.find(known_role, start)
                if pos == -1:
                    break
                
                # Check word boundaries
                before_ok = pos == 0 or not text[pos-1].isalnum()
                after_pos = pos + len(known_role)
                after_ok = after_pos >= len(text) or not text[after_pos].isalnum()
                
                if before_ok and after_ok:
                    # Get the actual case from original text
                    actual_role = text[pos:after_pos]
                    canonical, variant = self._get_canonical_role(actual_role)
                    
                    # v3.0.91d: Skip if canonical form is in false_positives
                    if canonical.lower() in self.false_positives:
                        start = pos + 1
                        continue
                    
                    match_key = (canonical, pos)
                    if match_key not in seen_matches:
                        seen_matches.add(match_key)
                        
                        # Get sentence context
                        context = self._get_sentence_context(text, pos, after_pos)
                        responsibility, action_type = self._extract_responsibility(context, actual_role)
                        
                        occurrence = RoleOccurrence(
                            role=actual_role,
                            context=context,
                            responsibility=responsibility,
                            action_type=action_type,
                            location=source_location,
                            confidence=0.90  # High confidence for known roles
                        )
                        
                        if canonical not in existing_roles:
                            # v3.0.12: Use helper to ensure entity_kind is populated
                            # v3.0.12b: Fixed NameError - use actual_role not match.group(0)
                            existing_roles[canonical] = self._create_extracted_role(canonical, actual_role)
                        
                        role_entry = existing_roles[canonical]
                        role_entry.variants.add(variant)
                        role_entry.occurrences.append(occurrence)
                        
                        if responsibility:
                            role_entry.responsibilities.append(responsibility)
                        role_entry.action_types[action_type] += 1
                
                start = pos + 1
        
        return existing_roles
    
    def extract_from_docx(self, filepath: str) -> Dict[str, ExtractedRole]:
        """Extract roles from a Word document."""
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx required. Install: pip install python-docx")
        
        doc = Document(filepath)
        all_roles: Dict[str, ExtractedRole] = {}
        
        for para_num, paragraph in enumerate(doc.paragraphs, 1):
            text = paragraph.text.strip()
            if text:
                para_roles = self.extract_from_text(text, f"Paragraph {para_num}")
                self._merge_roles(all_roles, para_roles)
        
        for table_num, table in enumerate(doc.tables, 1):
            for row_num, row in enumerate(table.rows, 1):
                for cell_num, cell in enumerate(row.cells, 1):
                    text = cell.text.strip()
                    if text:
                        cell_roles = self.extract_from_text(
                            text, f"Table {table_num}, Row {row_num}, Cell {cell_num}"
                        )
                        self._merge_roles(all_roles, cell_roles)
        
        return all_roles
    
    def extract_from_pdf(self, filepath: str) -> Dict[str, ExtractedRole]:
        """Extract roles from a PDF document with enhanced table support."""
        try:
            import pdfplumber
            use_pdfplumber = True
        except ImportError:
            use_pdfplumber = False
            try:
                import PyPDF2
            except ImportError:
                raise ImportError("pdfplumber or PyPDF2 required. Install: pip install pdfplumber")
        
        all_roles: Dict[str, ExtractedRole] = {}
        
        # First, try enhanced table extraction for better accuracy
        table_roles = self._extract_roles_from_pdf_tables(filepath)
        self._merge_roles(all_roles, table_roles)
        _log(f"Found {len(table_roles)} roles from table extraction", level='debug')
        
        # Then extract from text
        if use_pdfplumber:
            with pdfplumber.open(filepath) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    text = page.extract_text() or ""
                    if text:
                        page_roles = self.extract_from_text(text, f"Page {page_num}")
                        self._merge_roles(all_roles, page_roles)
        else:
            with open(filepath, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                for page_num, page in enumerate(reader.pages, 1):
                    text = page.extract_text() or ""
                    if text:
                        page_roles = self.extract_from_text(text, f"Page {page_num}")
                        self._merge_roles(all_roles, page_roles)
        
        return all_roles
    
    def _extract_roles_from_pdf_tables(self, filepath: str) -> Dict[str, ExtractedRole]:
        """
        Extract roles specifically from PDF tables with enhanced confidence.
        
        Tables often contain RACI matrices, responsibility assignments, and
        role definitions with higher reliability than free text.
        """
        roles: Dict[str, ExtractedRole] = {}
        
        try:
            from enhanced_table_extractor import EnhancedTableExtractor
            extractor = EnhancedTableExtractor(prefer_accuracy=True)
            result = extractor.extract_tables(filepath)
            
            for table in result.tables:
                # Check if this looks like a RACI matrix or responsibility table
                is_raci = self._is_raci_table(table.headers, table.rows)
                is_responsibility_table = self._is_responsibility_table(table.headers)
                
                # Confidence boost for table-sourced roles
                confidence_boost = 0.20 if is_raci else (0.15 if is_responsibility_table else 0.10)
                
                # Extract from headers (often contain role names)
                for header in table.headers:
                    header_roles = self.extract_from_text(
                        header, f"Table {table.index} Header"
                    )
                    for canonical, role_data in header_roles.items():
                        # Boost confidence for table headers
                        for occ in role_data.occurrences:
                            occ.confidence = min(1.0, occ.confidence + confidence_boost)
                        self._merge_role(roles, canonical, role_data)
                
                # Extract from table cells
                for row_idx, row in enumerate(table.rows):
                    for col_idx, cell in enumerate(row):
                        if cell and len(cell.strip()) > 2:
                            cell_roles = self.extract_from_text(
                                cell, f"Table {table.index}, Row {row_idx+1}"
                            )
                            for canonical, role_data in cell_roles.items():
                                for occ in role_data.occurrences:
                                    occ.confidence = min(1.0, occ.confidence + confidence_boost * 0.75)
                                self._merge_role(roles, canonical, role_data)
                                
        except ImportError:
            _log("Enhanced table extractor not available, using basic extraction", level='debug')
        except Exception as e:
            _log(f"Table role extraction failed: {e}", level='debug')
        
        return roles
    
    def _is_raci_table(self, headers: List[str], rows: List[List[str]]) -> bool:
        """Detect if a table is a RACI matrix."""
        # Check headers for RACI indicators
        raci_headers = {'r', 'a', 'c', 'i', 'responsible', 'accountable', 'consulted', 'informed'}
        header_text = ' '.join(h.lower() for h in headers)
        
        if any(raci in header_text for raci in ['raci', 'rasci', 'responsibility matrix']):
            return True
        
        # Check if cells contain mostly R/A/C/I values
        raci_values = {'r', 'a', 'c', 'i', 'x', '●', '✓', ''}
        raci_count = 0
        total_cells = 0
        
        for row in rows[:5]:  # Check first few rows
            for cell in row[1:]:  # Skip first column (usually role names)
                total_cells += 1
                if cell.strip().lower() in raci_values or len(cell.strip()) <= 2:
                    raci_count += 1
        
        return total_cells > 0 and raci_count / total_cells > 0.6
    
    def _is_responsibility_table(self, headers: List[str]) -> bool:
        """Detect if a table contains responsibility assignments."""
        responsibility_keywords = {
            'role', 'responsibility', 'function', 'task', 'activity',
            'action', 'owner', 'lead', 'support', 'department', 'organization'
        }
        header_text = ' '.join(h.lower() for h in headers)
        return any(kw in header_text for kw in responsibility_keywords)
    
    def _merge_role(self, target: Dict[str, ExtractedRole], 
                    canonical: str, role_data: ExtractedRole):
        """Merge a single role into target dictionary."""
        if canonical not in target:
            target[canonical] = role_data
        else:
            existing = target[canonical]
            existing.variants.update(role_data.variants)
            existing.occurrences.extend(role_data.occurrences)
            existing.responsibilities.extend(role_data.responsibilities)
    
    def _merge_roles(self, target: Dict[str, ExtractedRole], source: Dict[str, ExtractedRole]):
        """Merge source roles into target."""
        for canonical, role_data in source.items():
            if canonical not in target:
                target[canonical] = role_data
            else:
                existing = target[canonical]
                existing.variants.update(role_data.variants)
                existing.occurrences.extend(role_data.occurrences)
                existing.responsibilities.extend(role_data.responsibilities)
                for action, count in role_data.action_types.items():
                    existing.action_types[action] += count
    
    def generate_report(self, roles: Dict[str, ExtractedRole], 
                       min_confidence: float = 0.5,
                       min_occurrences: int = 1,
                       sort_by: str = 'frequency') -> str:
        """Generate a formatted report of extracted roles."""
        lines = []
        lines.append("=" * 80)
        lines.append("ROLE EXTRACTION REPORT")
        lines.append("=" * 80)
        lines.append("")
        
        filtered_roles = [
            (name, data) for name, data in roles.items()
            if data.avg_confidence >= min_confidence and data.frequency >= min_occurrences
        ]
        
        if sort_by == 'frequency':
            filtered_roles.sort(key=lambda x: (-x[1].frequency, -x[1].avg_confidence, x[0]))
        elif sort_by == 'confidence':
            filtered_roles.sort(key=lambda x: (-x[1].avg_confidence, -x[1].frequency, x[0]))
        else:
            filtered_roles.sort(key=lambda x: x[0])
        
        lines.append(f"Total unique roles found: {len(filtered_roles)}")
        lines.append(f"Filters: min_confidence={min_confidence}, min_occurrences={min_occurrences}")
        lines.append("")
        lines.append("-" * 80)
        
        for name, data in filtered_roles:
            lines.append(f"\n■ {name}")
            lines.append(f"  Frequency: {data.frequency} | Confidence: {data.avg_confidence:.0%}")
            
            if data.variants and len(data.variants) > 1:
                variants = sorted(v for v in data.variants if v.lower() != name.lower())
                if variants:
                    lines.append(f"  Also appears as: {', '.join(variants[:5])}")
            
            actions = {k: v for k, v in data.action_types.items() if k != 'unknown'}
            if actions:
                action_str = ', '.join(f"{k}({v})" for k, v in sorted(actions.items(), key=lambda x: -x[1])[:5])
                lines.append(f"  Actions: {action_str}")
            
            if data.responsibilities:
                unique_resp = list(set(data.responsibilities))[:3]
                lines.append(f"  Responsibilities:")
                for resp in unique_resp:
                    lines.append(f"    • {resp[:80]}{'...' if len(resp) > 80 else ''}")
            
            if data.occurrences:
                occ = data.occurrences[0]
                ctx = occ.context[:120] + ('...' if len(occ.context) > 120 else '')
                lines.append(f"  Example: \"{ctx}\"")
                lines.append(f"           [{occ.location}]")
        
        lines.append("")
        lines.append("=" * 80)
        lines.append("END OF REPORT")
        lines.append("=" * 80)
        
        return '\n'.join(lines)
    
    def export_to_csv(self, roles: Dict[str, ExtractedRole], filepath: str,
                     min_confidence: float = 0.5, min_occurrences: int = 1):
        """Export extracted roles to CSV format."""
        filtered_roles = [
            (name, data) for name, data in roles.items()
            if data.avg_confidence >= min_confidence and data.frequency >= min_occurrences
        ]
        filtered_roles.sort(key=lambda x: (-x[1].frequency, -x[1].avg_confidence))
        
        with open(filepath, 'w', newline='', encoding='utf-8') as f:
            writer = csv.writer(f)
            writer.writerow([
                'Role', 'Frequency', 'Confidence', 'Variants', 
                'Primary_Actions', 'Sample_Responsibilities', 'Source_Locations'
            ])
            
            for name, data in filtered_roles:
                actions = {k: v for k, v in data.action_types.items() if k != 'unknown'}
                writer.writerow([
                    name,
                    data.frequency,
                    f"{data.avg_confidence:.0%}",
                    '; '.join(sorted(data.variants)),
                    '; '.join(f"{k}({v})" for k, v in sorted(actions.items(), key=lambda x: -x[1])[:3]),
                    '; '.join(list(set(data.responsibilities))[:3]),
                    '; '.join(sorted(set(occ.location for occ in data.occurrences[:5])))
                ])
        
        _log(f"Exported {len(filtered_roles)} roles to {filepath}")


# =============================================================================
# COMPREHENSIVE TESTING
# =============================================================================

def run_tests():
    """Run comprehensive tests against NASA document text."""
    
    # Extended test corpus from NASA Systems Engineering Handbook
    test_corpus = """
    The systems engineer is skilled in the art and science of balancing organizational, cost, and 
    technical interactions in complex systems. The systems engineer and supporting organization are 
    vital to supporting program and Project Planning and Control (PP&C) with accurate and timely cost 
    and schedule information for the technical activities.
    
    The lead systems engineer ensures that the system technically fulfills the defined needs and 
    requirements and that a proper systems engineering approach is being followed. The systems engineer 
    oversees the project's systems engineering activities as performed by the technical team and directs, 
    communicates, monitors, and coordinates tasks. The systems engineer reviews and evaluates the 
    technical aspects of the project.
    
    The project manager has overall responsibility for managing the project team and ensuring that 
    the project delivers a technically correct system within cost and schedule. The project manager 
    may sometimes perform these practices for small projects.
    
    Systems engineering plays a key role in the project organization. Managing a project consists of 
    three main objectives: managing the technical aspects of the project, managing the project team, 
    and managing the cost and schedule.
    
    The exact role and responsibility of the systems engineer may change from project to project 
    depending on the size and complexity of the project.
    
    For large projects, there may be one or more systems engineers. The systems engineer usually 
    plays the key role in leading the development of the concept of operations (ConOps).
    
    The Configuration Manager maintains the integrity and traceability of product configuration 
    throughout the system life cycle.
    
    Technical reviews are conducted by the technical team and reviewed by the Chief Engineer.
    
    The Mission Directorate Associate Administrator (MDAA) approves major program decisions.
    
    Integration activities are coordinated with the Integration Lead and verified by Quality Assurance.
    
    The Principal Investigator is responsible for the scientific objectives of the mission.
    
    Safety reviews are conducted by the Safety Engineer and approved by the System Safety Panel.
    
    The Contracting Officer Representative (COR) monitors contractor performance and ensures 
    compliance with contract requirements.
    
    The Risk Manager is responsible for identifying, analyzing, and mitigating project risks.
    
    Interface specifications are developed by the Interface Control Working Group (ICWG) and approved 
    by the Configuration Control Board (CCB).
    
    The Software Lead coordinates with the Hardware Lead to ensure system integration.
    
    Verification activities are performed by the Test Engineer and validated by Quality Assurance.
    
    The Technical Authority provides independent technical oversight of the project.
    
    Requirements are traced by the Requirements Manager and verified by the Verification Lead.
    
    The Flight Director is responsible for all real-time mission operations.
    
    Program-level decisions require approval from the Program Manager and the Chief Engineer.
    
    The Integration Engineer assembles components according to the integration plan and works with 
    the Test Engineer to verify proper assembly.
    
    Mission assurance is provided by the Mission Assurance Manager who ensures compliance with 
    quality standards.
    
    The Deputy Project Manager assists the Project Manager in managing day-to-day operations.
    
    The Reliability Engineer analyzes system reliability and provides recommendations for 
    improvement.
    
    Configuration audits are conducted by the Configuration Manager with support from Quality 
    Assurance.
    
    The technical team supports the systems engineer in performing technical assessments.
    """
    
    # Expected roles that should be found
    expected_roles = [
        'Systems Engineer', 'Lead Systems Engineer', 'Project Manager', 'Chief Engineer',
        'Configuration Manager', 'Principal Investigator', 'Safety Engineer',
        'Risk Manager', 'Software Lead', 'Hardware Lead', 'Test Engineer',
        'Technical Authority', 'Requirements Manager', 'Verification Lead',
        'Flight Director', 'Program Manager', 'Integration Lead', 'Integration Engineer',
        'Mission Directorate Associate Administrator', 'Contracting Officer Representative',
        'Interface Control Working Group', 'Configuration Control Board', 'System Safety Panel',
        'Quality Assurance', 'Technical Team', 'Mission Assurance Manager',
        'Deputy Project Manager', 'Reliability Engineer'
    ]
    
    print("=" * 80)
    print("ROLE EXTRACTOR v3.0 - COMPREHENSIVE TEST")
    print("=" * 80)
    print()
    
    extractor = RoleExtractor()
    roles = extractor.extract_from_text(test_corpus, "Test Corpus")
    
    # Generate report
    report = extractor.generate_report(roles, min_confidence=0.5, min_occurrences=1)
    print(report)
    
    # Validation
    print("\n" + "=" * 80)
    print("VALIDATION RESULTS")
    print("=" * 80)
    
    found_roles = set(roles.keys())
    expected_set = set(expected_roles)
    
    # Check coverage
    matched = found_roles & expected_set
    missed = expected_set - found_roles
    extra = found_roles - expected_set
    
    print(f"\nExpected roles: {len(expected_set)}")
    print(f"Found roles: {len(found_roles)}")
    print(f"Matched: {len(matched)} ({len(matched)/len(expected_set)*100:.1f}%)")
    
    if missed:
        print(f"\nMissed ({len(missed)}):")
        for role in sorted(missed):
            print(f"  - {role}")
    
    if extra:
        print(f"\nAdditional roles found ({len(extra)}):")
        for role in sorted(extra):
            data = roles[role]
            print(f"  + {role} (freq={data.frequency}, conf={data.avg_confidence:.0%})")
    
    # Calculate accuracy metrics
    precision = len(matched) / len(found_roles) if found_roles else 0
    recall = len(matched) / len(expected_set) if expected_set else 0
    f1 = 2 * precision * recall / (precision + recall) if (precision + recall) > 0 else 0
    
    print(f"\n--- Metrics ---")
    print(f"Precision: {precision:.1%}")
    print(f"Recall: {recall:.1%}")
    print(f"F1 Score: {f1:.1%}")
    
    return roles


if __name__ == "__main__":
    run_tests()
