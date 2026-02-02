"""
Hyperlink Validator Export Functions
====================================
Export validation results to various formats: CSV, JSON, HTML, and highlighted
copies of original documents (DOCX, Excel).

This module is designed to be independent and can be tested separately.

Exclusion Handling:
- URLs marked as excluded with treat_as_valid=True are shown as WORKING/OK
- The exclusion_reason field indicates why the URL was excluded

Highlighted Export (v3.0.110):
- export_highlighted_docx: Creates DOCX with broken links highlighted in red
- export_highlighted_excel: Creates Excel with rows containing broken links in red
"""

import csv
import json
import io
import os
import re
import zipfile
import copy
import xml.etree.ElementTree as ET
from datetime import datetime
from typing import List, Optional, Any, Dict, Set, Tuple

from .models import ValidationResult, ValidationSummary, ValidationRun

# Check for python-docx availability
try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_COLOR_INDEX
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Check for openpyxl availability
try:
    from openpyxl import load_workbook
    from openpyxl.styles import PatternFill, Font, Border, Side
    from openpyxl.comments import Comment
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False


def _apply_exclusion_display(result: ValidationResult) -> ValidationResult:
    """
    Apply exclusion display rules: excluded URLs with treat_as_valid=True
    should show as WORKING in exports.

    Returns a copy of the result with adjusted status if applicable.
    """
    if result.excluded and getattr(result, 'treat_as_valid', True):
        # Create a modified copy for display
        modified = ValidationResult(
            url=result.url,
            status='WORKING',  # Show as OK
            status_code=200,  # Synthetic OK code
            message=f'Excluded: {result.exclusion_reason}' if result.exclusion_reason else 'Excluded (treated as valid)',
            redirect_url=result.redirect_url,
            redirect_count=result.redirect_count,
            response_time_ms=0,
            dns_resolved=True,
            ssl_valid=True,
            auth_used=result.auth_used,
            attempts=0,
            checked_at=result.checked_at,
            # Extended fields
            dns_ip_addresses=result.dns_ip_addresses,
            ssl_issuer=result.ssl_issuer,
            ssl_days_until_expiry=result.ssl_days_until_expiry,
            is_soft_404=False,
            is_suspicious=False,
            domain_category=result.domain_category,
            excluded=True,
            exclusion_reason=result.exclusion_reason
        )
        return modified
    return result


def export_csv(
    results: List[ValidationResult],
    summary: Optional[ValidationSummary] = None,
    apply_exclusion_rules: bool = True
) -> str:
    """
    Export results to CSV format.

    Args:
        results: List of validation results
        summary: Optional summary statistics
        apply_exclusion_rules: If True, excluded URLs show as WORKING

    Returns:
        CSV content as string
    """
    output = io.StringIO()

    # Use UTF-8 BOM for Excel compatibility
    output.write('\ufeff')

    writer = csv.writer(output, quoting=csv.QUOTE_MINIMAL)

    # Header row
    writer.writerow([
        'URL',
        'Status',
        'Status Code',
        'Message',
        'Redirect URL',
        'Redirect Count',
        'Response Time (ms)',
        'DNS Resolved',
        'SSL Valid',
        'Auth Used',
        'Attempts',
        'Excluded',
        'Checked At'
    ])

    # Data rows
    for result in results:
        # Apply exclusion display rules
        display_result = _apply_exclusion_display(result) if apply_exclusion_rules else result

        writer.writerow([
            display_result.url,
            display_result.status,
            display_result.status_code or '',
            display_result.message,
            display_result.redirect_url or '',
            display_result.redirect_count,
            f'{display_result.response_time_ms:.1f}' if display_result.response_time_ms else '',
            'Yes' if display_result.dns_resolved else 'No',
            'Yes' if display_result.ssl_valid else 'No',
            display_result.auth_used,
            display_result.attempts,
            'Yes' if display_result.excluded else 'No',
            display_result.checked_at
        ])

    # Add summary section
    if summary:
        writer.writerow([])  # Empty row
        writer.writerow(['SUMMARY'])
        writer.writerow(['Total URLs', summary.total])
        writer.writerow(['Working', summary.working])
        writer.writerow(['Broken', summary.broken])
        writer.writerow(['Redirect', summary.redirect])
        writer.writerow(['Timeout', summary.timeout])
        writer.writerow(['Blocked', summary.blocked])
        writer.writerow(['DNS Failed', summary.dns_failed])
        writer.writerow(['SSL Error', summary.ssl_error])
        writer.writerow(['Invalid', summary.invalid])
        writer.writerow(['Unknown', summary.unknown])
        writer.writerow(['Success Rate', f'{summary.success_rate:.1f}%'])
        writer.writerow(['Average Response', f'{summary.average_response_ms:.1f}ms'])
        writer.writerow(['Total Time', f'{summary.total_time_seconds:.1f}s'])

    return output.getvalue()


def export_json(
    results: List[ValidationResult],
    summary: Optional[ValidationSummary] = None,
    run: Optional[ValidationRun] = None,
    apply_exclusion_rules: bool = True
) -> str:
    """
    Export results to JSON format.

    Args:
        results: List of validation results
        summary: Optional summary statistics
        run: Optional run metadata
        apply_exclusion_rules: If True, excluded URLs show as WORKING

    Returns:
        JSON content as string
    """
    # Apply exclusion display rules
    display_results = [
        _apply_exclusion_display(r) if apply_exclusion_rules else r
        for r in results
    ]

    data = {
        'exported_at': datetime.utcnow().isoformat() + 'Z',
        'exporter': 'TechWriterReview HyperlinkValidator v1.0.0',
        'results': [r.to_dict() for r in display_results]
    }

    if summary:
        data['summary'] = summary.to_dict()

    if run:
        data['run'] = {
            'run_id': run.run_id,
            'job_id': run.job_id,
            'created_at': run.created_at,
            'completed_at': run.completed_at,
            'mode': run.mode,
            'status': run.status
        }

    return json.dumps(data, indent=2)


def export_html(
    results: List[ValidationResult],
    summary: Optional[ValidationSummary] = None,
    run: Optional[ValidationRun] = None,
    apply_exclusion_rules: bool = True
) -> str:
    """
    Export results to HTML report format.

    Args:
        results: List of validation results
        summary: Optional summary statistics
        run: Optional run metadata
        apply_exclusion_rules: If True, excluded URLs show as WORKING

    Returns:
        HTML content as string
    """
    # Apply exclusion display rules
    display_results = [
        _apply_exclusion_display(r) if apply_exclusion_rules else r
        for r in results
    ]

    # Get status color class
    def status_class(status: str, excluded: bool = False) -> str:
        if excluded:
            return 'status-excluded'  # Special styling for excluded items
        status_map = {
            'WORKING': 'status-working',
            'REDIRECT': 'status-redirect',
            'BROKEN': 'status-broken',
            'TIMEOUT': 'status-timeout',
            'BLOCKED': 'status-blocked',
            'DNSFAILED': 'status-dns',
            'SSLERROR': 'status-ssl',
            'INVALID': 'status-invalid',
            'UNKNOWN': 'status-unknown',
            'PENDING': 'status-pending',
            'SKIPPED': 'status-skipped'
        }
        return status_map.get(status.upper(), 'status-unknown')

    # Build HTML
    html = f'''<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Hyperlink Validation Report</title>
    <style>
        :root {{
            --color-working: #22c55e;
            --color-redirect: #3b82f6;
            --color-broken: #ef4444;
            --color-timeout: #f59e0b;
            --color-blocked: #8b5cf6;
            --color-dns: #ec4899;
            --color-ssl: #f97316;
            --color-invalid: #6b7280;
            --color-unknown: #9ca3af;
        }}

        * {{
            box-sizing: border-box;
            margin: 0;
            padding: 0;
        }}

        body {{
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, 'Helvetica Neue', Arial, sans-serif;
            line-height: 1.6;
            color: #1f2937;
            background: #f9fafb;
            padding: 2rem;
        }}

        .container {{
            max-width: 1400px;
            margin: 0 auto;
        }}

        h1 {{
            font-size: 1.875rem;
            font-weight: 700;
            margin-bottom: 0.5rem;
            color: #111827;
        }}

        .subtitle {{
            color: #6b7280;
            margin-bottom: 2rem;
        }}

        .summary-grid {{
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(150px, 1fr));
            gap: 1rem;
            margin-bottom: 2rem;
        }}

        .summary-card {{
            background: white;
            border-radius: 8px;
            padding: 1rem;
            box-shadow: 0 1px 3px rgba(0, 0, 0, 0.1);
            text-align: center;
        }}

        .summary-card .value {{
            font-size: 2rem;
            font-weight: 700;
            margin-bottom: 0.25rem;
        }}

        .summary-card .label {{
            font-size: 0.875rem;
            color: #6b7280;
        }}

        .summary-card.working .value {{ color: var(--color-working); }}
        .summary-card.broken .value {{ color: var(--color-broken); }}
        .summary-card.redirect .value {{ color: var(--color-redirect); }}
        .summary-card.timeout .value {{ color: var(--color-timeout); }}
        .summary-card.blocked .value {{ color: var(--color-blocked); }}

        .results-table {{
            width: 100%;
            background: white;
            border-radius: 8px;
            box-shadow: 0 1px 3px rgba(0, 0, 0, 0.1);
            overflow: hidden;
        }}

        .results-table table {{
            width: 100%;
            border-collapse: collapse;
        }}

        .results-table th {{
            background: #f3f4f6;
            padding: 0.75rem 1rem;
            text-align: left;
            font-weight: 600;
            font-size: 0.875rem;
            color: #374151;
            border-bottom: 1px solid #e5e7eb;
        }}

        .results-table td {{
            padding: 0.75rem 1rem;
            border-bottom: 1px solid #f3f4f6;
            font-size: 0.875rem;
        }}

        .results-table tr:hover {{
            background: #f9fafb;
        }}

        .results-table .url {{
            max-width: 400px;
            word-break: break-all;
        }}

        .results-table .url a {{
            color: #2563eb;
            text-decoration: none;
        }}

        .results-table .url a:hover {{
            text-decoration: underline;
        }}

        .status-badge {{
            display: inline-flex;
            align-items: center;
            padding: 0.25rem 0.75rem;
            border-radius: 9999px;
            font-size: 0.75rem;
            font-weight: 600;
            text-transform: uppercase;
        }}

        .status-working {{ background: #dcfce7; color: #166534; }}
        .status-redirect {{ background: #dbeafe; color: #1e40af; }}
        .status-broken {{ background: #fee2e2; color: #991b1b; }}
        .status-timeout {{ background: #fef3c7; color: #92400e; }}
        .status-blocked {{ background: #ede9fe; color: #5b21b6; }}
        .status-dns {{ background: #fce7f3; color: #9d174d; }}
        .status-ssl {{ background: #ffedd5; color: #c2410c; }}
        .status-invalid {{ background: #f3f4f6; color: #374151; }}
        .status-unknown {{ background: #f3f4f6; color: #6b7280; }}
        .status-pending {{ background: #e0e7ff; color: #3730a3; }}
        .status-skipped {{ background: #f3f4f6; color: #9ca3af; }}
        .status-excluded {{ background: #d1fae5; color: #065f46; border: 1px dashed #10b981; }}

        .meta-info {{
            margin-top: 2rem;
            padding: 1rem;
            background: #f3f4f6;
            border-radius: 8px;
            font-size: 0.75rem;
            color: #6b7280;
        }}

        @media print {{
            body {{
                background: white;
                padding: 0;
            }}

            .results-table {{
                box-shadow: none;
                border: 1px solid #e5e7eb;
            }}
        }}
    </style>
</head>
<body>
    <div class="container">
        <h1>Hyperlink Validation Report</h1>
        <p class="subtitle">Generated: {datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S')} UTC</p>
'''

    # Summary section
    if summary:
        html += f'''
        <div class="summary-grid">
            <div class="summary-card">
                <div class="value">{summary.total}</div>
                <div class="label">Total URLs</div>
            </div>
            <div class="summary-card working">
                <div class="value">{summary.working}</div>
                <div class="label">Working</div>
            </div>
            <div class="summary-card broken">
                <div class="value">{summary.broken}</div>
                <div class="label">Broken</div>
            </div>
            <div class="summary-card redirect">
                <div class="value">{summary.redirect}</div>
                <div class="label">Redirect</div>
            </div>
            <div class="summary-card timeout">
                <div class="value">{summary.timeout}</div>
                <div class="label">Timeout</div>
            </div>
            <div class="summary-card blocked">
                <div class="value">{summary.blocked + summary.dns_failed + summary.ssl_error}</div>
                <div class="label">Other Errors</div>
            </div>
        </div>
'''

    # Results table
    html += '''
        <div class="results-table">
            <table>
                <thead>
                    <tr>
                        <th>Status</th>
                        <th>URL</th>
                        <th>Code</th>
                        <th>Message</th>
                        <th>Time</th>
                    </tr>
                </thead>
                <tbody>
'''

    for result in display_results:
        time_str = f'{result.response_time_ms:.0f}ms' if result.response_time_ms else '-'
        code_str = str(result.status_code) if result.status_code else '-'
        is_excluded = getattr(result, 'excluded', False)

        html += f'''
                    <tr>
                        <td><span class="status-badge {status_class(result.status, is_excluded)}">{result.status}{' (Excluded)' if is_excluded else ''}</span></td>
                        <td class="url"><a href="{result.url}" target="_blank" rel="noopener">{result.url}</a></td>
                        <td>{code_str}</td>
                        <td>{result.message}</td>
                        <td>{time_str}</td>
                    </tr>
'''

    html += '''
                </tbody>
            </table>
        </div>
'''

    # Meta info
    if run:
        html += f'''
        <div class="meta-info">
            <strong>Run ID:</strong> {run.run_id} |
            <strong>Mode:</strong> {run.mode} |
            <strong>Started:</strong> {run.created_at} |
            <strong>Completed:</strong> {run.completed_at or 'N/A'}
'''
        if summary:
            html += f''' |
            <strong>Total Time:</strong> {summary.total_time_seconds:.1f}s |
            <strong>Success Rate:</strong> {summary.success_rate:.1f}%
'''
        html += '''
        </div>
'''

    html += '''
    </div>
</body>
</html>
'''

    return html


# =============================================================================
# HIGHLIGHTED DOCUMENT EXPORTS (v3.0.110)
# =============================================================================

def _get_broken_urls(results: List[ValidationResult]) -> Set[str]:
    """
    Get set of URLs that are broken/failed validation.

    Includes: BROKEN, TIMEOUT, DNSFAILED, SSLERROR, INVALID statuses.
    Excludes: WORKING, REDIRECT (these are OK), excluded URLs treated as valid.
    """
    broken_statuses = {'BROKEN', 'TIMEOUT', 'DNSFAILED', 'SSLERROR', 'INVALID', 'BLOCKED'}
    broken_urls = set()

    for result in results:
        # Skip excluded URLs that are treated as valid
        if result.excluded and getattr(result, 'treat_as_valid', True):
            continue

        if result.status.upper() in broken_statuses:
            broken_urls.add(result.url)
            # Also add normalized versions (without trailing slash, etc.)
            normalized = result.url.rstrip('/')
            broken_urls.add(normalized)
            if result.url.startswith('http://'):
                broken_urls.add(result.url.replace('http://', 'https://'))
            elif result.url.startswith('https://'):
                broken_urls.add(result.url.replace('https://', 'http://'))

    return broken_urls


def _get_result_for_url(url: str, results: List[ValidationResult]) -> Optional[ValidationResult]:
    """Get the validation result for a specific URL."""
    for result in results:
        if result.url == url or result.url.rstrip('/') == url.rstrip('/'):
            return result
    return None


def export_highlighted_docx(
    source_path: str,
    results: List[ValidationResult],
    output_path: Optional[str] = None
) -> Tuple[bool, str, bytes]:
    """
    Create a copy of a DOCX file with broken hyperlinks highlighted.

    Broken links are marked with:
    - Red text color
    - Yellow highlight background
    - Strikethrough formatting
    - A comment indicating the error

    Args:
        source_path: Path to the original DOCX file
        results: List of validation results
        output_path: Optional path for output file (if None, returns bytes)

    Returns:
        Tuple of (success: bool, message: str, file_bytes: bytes)
    """
    if not DOCX_AVAILABLE:
        return False, "python-docx library not installed. Cannot create highlighted DOCX.", b''

    if not os.path.exists(source_path):
        return False, f"Source file not found: {source_path}", b''

    broken_urls = _get_broken_urls(results)

    if not broken_urls:
        return False, "No broken links found to highlight.", b''

    try:
        # Load the document
        doc = Document(source_path)
        highlighted_count = 0

        # Process all paragraphs
        for para in doc.paragraphs:
            highlighted_count += _highlight_broken_links_in_paragraph(para, broken_urls, results)

        # Process tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        highlighted_count += _highlight_broken_links_in_paragraph(para, broken_urls, results)

        # Process headers and footers
        for section in doc.sections:
            for header in [section.header, section.first_page_header, section.even_page_header]:
                if header:
                    for para in header.paragraphs:
                        highlighted_count += _highlight_broken_links_in_paragraph(para, broken_urls, results)
            for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
                if footer:
                    for para in footer.paragraphs:
                        highlighted_count += _highlight_broken_links_in_paragraph(para, broken_urls, results)

        # Save to bytes buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        file_bytes = buffer.getvalue()

        # Optionally save to file
        if output_path:
            with open(output_path, 'wb') as f:
                f.write(file_bytes)

        return True, f"Highlighted {highlighted_count} broken link(s) in document.", file_bytes

    except Exception as e:
        return False, f"Error processing DOCX: {str(e)}", b''


def _highlight_broken_links_in_paragraph(
    para,
    broken_urls: Set[str],
    results: List[ValidationResult]
) -> int:
    """
    Highlight broken hyperlinks in a paragraph.

    Returns the number of links highlighted.
    """
    highlighted = 0

    # Access the underlying XML to find hyperlinks
    try:
        # Get hyperlink elements from paragraph XML
        para_xml = para._element

        # Find all hyperlink elements
        hyperlinks = para_xml.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hyperlink')

        for hyperlink in hyperlinks:
            # Get the relationship ID
            r_id = hyperlink.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')

            if r_id:
                # Try to get the URL from relationships
                try:
                    # Access document part to get relationships
                    part = para.part
                    if hasattr(part, 'rels') and r_id in part.rels:
                        rel = part.rels[r_id]
                        url = rel.target_ref if hasattr(rel, 'target_ref') else str(rel._target)

                        # Check if this URL is broken
                        if url in broken_urls or url.rstrip('/') in broken_urls:
                            # Highlight all runs within this hyperlink
                            runs = hyperlink.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r')
                            for run in runs:
                                _apply_broken_link_formatting(run)

                            highlighted += 1
                except Exception:
                    pass

        # Also check for URLs in plain text (not hyperlinked)
        for run in para.runs:
            if run.text:
                for url in broken_urls:
                    if url in run.text:
                        _apply_broken_link_formatting_to_run(run)
                        highlighted += 1
                        break

    except Exception:
        pass

    return highlighted


def _apply_broken_link_formatting(run_element):
    """Apply red/strikethrough formatting to a run XML element."""
    try:
        # Get or create run properties
        rPr = run_element.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr')
        if rPr is None:
            rPr = OxmlElement('w:rPr')
            run_element.insert(0, rPr)

        # Add red color
        color = OxmlElement('w:color')
        color.set(qn('w:val'), 'FF0000')  # Red
        rPr.append(color)

        # Add yellow highlight
        highlight = OxmlElement('w:highlight')
        highlight.set(qn('w:val'), 'yellow')
        rPr.append(highlight)

        # Add strikethrough
        strike = OxmlElement('w:strike')
        strike.set(qn('w:val'), 'true')
        rPr.append(strike)

    except Exception:
        pass


def _apply_broken_link_formatting_to_run(run):
    """Apply broken link formatting to a python-docx Run object."""
    try:
        run.font.color.rgb = RGBColor(255, 0, 0)  # Red
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        run.font.strike = True
    except Exception:
        pass


def export_highlighted_excel(
    source_path: str,
    results: List[ValidationResult],
    output_path: Optional[str] = None,
    link_column: Optional[int] = None
) -> Tuple[bool, str, bytes]:
    """
    Create a copy of an Excel file with rows containing broken links highlighted.

    Broken link rows are marked with:
    - Red background fill on the entire row
    - Bold red text on the URL cell
    - A comment on the URL cell with the error details

    Args:
        source_path: Path to the original Excel file
        results: List of validation results
        output_path: Optional path for output file (if None, returns bytes)
        link_column: Optional column index (1-based) containing URLs.
                     If None, will auto-detect.

    Returns:
        Tuple of (success: bool, message: str, file_bytes: bytes)
    """
    if not OPENPYXL_AVAILABLE:
        return False, "openpyxl library not installed. Cannot create highlighted Excel.", b''

    if not os.path.exists(source_path):
        return False, f"Source file not found: {source_path}", b''

    broken_urls = _get_broken_urls(results)

    if not broken_urls:
        return False, "No broken links found to highlight.", b''

    try:
        # Load workbook
        wb = load_workbook(source_path)
        highlighted_count = 0

        # Define styles for broken links
        red_fill = PatternFill(start_color='FFCCCC', end_color='FFCCCC', fill_type='solid')
        red_font = Font(color='CC0000', bold=True)
        error_fill = PatternFill(start_color='FF6666', end_color='FF6666', fill_type='solid')

        # Process each sheet
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]

            # Find URL columns if not specified
            url_columns = _find_url_columns(ws) if link_column is None else [link_column]

            if not url_columns:
                continue

            # Process each row
            for row_idx, row in enumerate(ws.iter_rows(min_row=2), start=2):  # Skip header
                row_has_broken_link = False
                broken_cells = []

                for col_idx in url_columns:
                    if col_idx <= len(row):
                        cell = row[col_idx - 1]
                        cell_value = str(cell.value) if cell.value else ''

                        # Check if cell contains a broken URL
                        for url in broken_urls:
                            if url in cell_value or cell_value in broken_urls:
                                row_has_broken_link = True
                                broken_cells.append((cell, url))
                                break

                        # Also check hyperlink target
                        if cell.hyperlink and cell.hyperlink.target:
                            target = cell.hyperlink.target
                            if target in broken_urls or target.rstrip('/') in broken_urls:
                                row_has_broken_link = True
                                broken_cells.append((cell, target))

                # Highlight the entire row if it has a broken link
                if row_has_broken_link:
                    highlighted_count += 1

                    # Apply red fill to entire row
                    for cell in row:
                        cell.fill = red_fill

                    # Apply stronger formatting to the URL cells
                    for cell, url in broken_cells:
                        cell.fill = error_fill
                        cell.font = red_font

                        # Add comment with error details
                        result = _get_result_for_url(url, results)
                        if result:
                            comment_text = f"BROKEN LINK\nStatus: {result.status}\nMessage: {result.message}"
                            if result.status_code:
                                comment_text += f"\nHTTP Code: {result.status_code}"
                            cell.comment = Comment(comment_text, "Hyperlink Validator")

        # Save to bytes buffer
        buffer = io.BytesIO()
        wb.save(buffer)
        buffer.seek(0)
        file_bytes = buffer.getvalue()

        # Optionally save to file
        if output_path:
            with open(output_path, 'wb') as f:
                f.write(file_bytes)

        wb.close()

        return True, f"Highlighted {highlighted_count} row(s) with broken links.", file_bytes

    except Exception as e:
        return False, f"Error processing Excel: {str(e)}", b''


def _find_url_columns(ws) -> List[int]:
    """
    Auto-detect columns that likely contain URLs.

    Looks for:
    - Header row with 'URL', 'Link', 'Hyperlink', 'Website', 'Web' in name
    - Columns with cells containing http://, https://, www., or mailto:

    Returns list of 1-based column indices.
    """
    url_columns = []
    url_patterns = ['url', 'link', 'hyperlink', 'website', 'web', 'href']
    url_regex = re.compile(r'https?://|www\.|mailto:', re.IGNORECASE)

    # Check header row
    header_row = list(ws.iter_rows(min_row=1, max_row=1))[0] if ws.max_row > 0 else []

    for col_idx, cell in enumerate(header_row, start=1):
        if cell.value:
            header_text = str(cell.value).lower()
            if any(pattern in header_text for pattern in url_patterns):
                url_columns.append(col_idx)

    # If no headers matched, scan first few rows for URL-like content
    if not url_columns:
        url_col_candidates = {}

        for row in ws.iter_rows(min_row=1, max_row=min(20, ws.max_row)):
            for col_idx, cell in enumerate(row, start=1):
                if cell.value:
                    cell_text = str(cell.value)
                    if url_regex.search(cell_text) or (cell.hyperlink and cell.hyperlink.target):
                        url_col_candidates[col_idx] = url_col_candidates.get(col_idx, 0) + 1

        # Select columns with at least 2 URL-like values
        url_columns = [col for col, count in url_col_candidates.items() if count >= 2]

    # If still nothing, check all columns with hyperlinks
    if not url_columns:
        for col_idx in range(1, ws.max_column + 1):
            for row in ws.iter_rows(min_row=2, max_row=min(100, ws.max_row)):
                cell = row[col_idx - 1] if col_idx <= len(row) else None
                if cell and cell.hyperlink:
                    url_columns.append(col_idx)
                    break

    return url_columns


def is_highlighted_export_available() -> Dict[str, bool]:
    """Check which highlighted export formats are available."""
    return {
        'docx': DOCX_AVAILABLE,
        'excel': OPENPYXL_AVAILABLE
    }
